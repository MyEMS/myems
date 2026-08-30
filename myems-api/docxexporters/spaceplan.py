"""
Space Plan DOCX Exporter

This module provides functionality to export space plan data to DOCX format.
It generates comprehensive reports showing energy plan analysis for spaces
with detailed breakdown by energy categories and time periods.

Key Features:
- Space energy plan analysis (plan - actual)
- Base period vs reporting period comparison
- Plan breakdown by energy categories (TCE/TCO2E)
- Detailed data with line charts
- Multi-language support
- Base64 encoding for file transmission

The exported DOCX file includes:
- Cover page with logo and report metadata (Space Data - Plan Analysis)
- Combined analysis page (plan summary table + TCE/TCO2E breakdown tables with pie charts)
- Detailed data charts (paginated, up to 4 per page in 2x2 grid)
- Parameter data tables + line charts with filled area
"""

import base64
import os
import time
import uuid
import io

from decimal import Decimal
from typing import Optional, Dict, List, Any, BinaryIO
import logging

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.utilities import get_translation, round2

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

_font_setup_done = False


def setup_chinese_fonts():
    global _font_setup_done
    if _font_setup_done:
        return True

    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             '..', 'pdfexporters', 'fonts', 'NotoSansCJK-Regular.ttc')
    try:
        import matplotlib.font_manager as fm
        fm.fontManager.addfont(font_path)
        prop = fm.FontProperties(fname=font_path)
        font_name = prop.get_name()
        plt.rcParams['font.sans-serif'] = [font_name]
        plt.rcParams['axes.unicode_minus'] = False
        logger.info(f"Successfully loaded bundled font: {font_name} from {font_path}")
        _font_setup_done = True
        return True
    except Exception as e:
        logger.warning(f"Failed to load bundled font from {font_path}: {e}")

    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    logger.warning("Failed to load bundled NotoSansCJK font, using DejaVu Sans")
    return False


def _convert_decimals(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    elif isinstance(obj, dict):
        return {k: _convert_decimals(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_convert_decimals(item) for item in obj]
    elif isinstance(obj, tuple):
        return tuple(_convert_decimals(item) for item in obj)
    return obj


def _set_min_row_height(cell, height_twips=150, exact=False):
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def _reset_cell_paragraph_spacing(cell, font_size=9):
    line_twips = max(int(font_size * 20 * 1.15), 120)
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(spacing)
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


def _set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'start', 'bottom', 'end']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _style_table_cell(cell, is_header=False, is_green=False, bold=False, font_size=9):
    cell.paragraphs[0].runs[0].font.bold = bold
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _reset_cell_paragraph_spacing(cell, font_size=font_size)
    _set_cell_margins_zero(cell)
    row_h = max(int(font_size * 20 * 1.3), 160)
    _set_min_row_height(cell, height_twips=row_h, exact=False)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    if is_header:
        shd.set(qn('w:fill'), 'D9E2F3')
    elif is_green:
        shd.set(qn('w:fill'), '90EE90')
    else:
        return
    tcPr.append(shd)


class SpacePlanDOCXExporter:
    """
    Export space plan data to DOCX format.
    Generates comprehensive reports with charts and tables matching Excel layout.
    """

    def __init__(self, language: str = 'zh_CN'):
        font_setup_success = setup_chinese_fonts()
        if not font_setup_success:
            logger.warning("Chinese font setup failed, some text may not display correctly")

        self.language = language
        self.trans = get_translation(language)
        self._ = self.trans.gettext

        self.dpi = 120

        self.chart_colors = ['#4472C4', '#ED7D31', '#70AD47', '#FFC000', '#5B9BD5',
                             '#FF6B6B', '#9B59B6', '#1ABC9C', '#E67E22', '#2ECC71',
                             '#3498DB', '#E74C3C', '#2ECC71', '#F39C12', '#9B59B6']

    def export(self,
               report: Dict[str, Any],
               name: str,
               base_period_start_datetime_local: str,
               base_period_end_datetime_local: str,
               reporting_start_datetime_local: str,
               reporting_end_datetime_local: str,
               period_type: str,
               language: str) -> Optional[str]:
        if report is None:
            return None
        start_time = time.time()
        logger.info(f"Starting DOCX generation for {name}")

        docx_filename = self.generate_docx(
            report, name,
            base_period_start_datetime_local,
            base_period_end_datetime_local,
            reporting_start_datetime_local,
            reporting_end_datetime_local,
            period_type,
            language
        )

        result = ''
        if docx_filename and os.path.exists(docx_filename):
            try:
                with open(docx_filename, 'rb') as binary_file:
                    binary_data = binary_file.read()
                result = base64.b64encode(binary_data).decode('utf-8')
            except Exception as e:
                logger.error(f"Failed to encode DOCX: {str(e)}")
            finally:
                try:
                    os.remove(docx_filename)
                except Exception:
                    pass
        elapsed = time.time() - start_time
        logger.info(f"DOCX generation completed in {elapsed:.2f}s for {name}")
        return result

    @staticmethod
    def _fig_to_bytesio(fig, dpi: int) -> BinaryIO:
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    def _make_pie_chart(self, values, labels, title, colors=None):
        if not values or sum((v or 0) for v in values) == 0:
            return None
        fig, ax = plt.subplots(figsize=(3.2, 2.6))
        if colors is None:
            colors = self.chart_colors[:len(labels)]
        filtered = [(l, v, c) for l, v, c in zip(labels, values, colors) if (v or 0) > 0]
        if not filtered:
            plt.close(fig)
            return None
        f_labels, f_values, f_colors = zip(*filtered)
        if len(f_labels) > 8:
            sorted_data = sorted(zip(f_labels, f_values, f_colors), key=lambda x: x[1], reverse=True)
            top_data = sorted_data[:7]
            other_sum = sum(v for _, v, _ in sorted_data[7:])
            f_labels = [l for l, _, _ in top_data] + [self._('Others')]
            f_values = [v for _, v, _ in top_data] + [other_sum]
            f_colors = [c for _, _, c in top_data] + ['#999999']
        ax.pie(f_values, labels=f_labels, autopct='%1.1f%%', colors=f_colors, startangle=90)
        ax.set_title(title, fontsize=10, fontweight='bold')
        return self._fig_to_bytesio(fig, self.dpi)

    def generate_docx(self,
                      report: Dict[str, Any],
                      name: str,
                      base_period_start_datetime_local: str,
                      base_period_end_datetime_local: str,
                      reporting_start_datetime_local: str,
                      reporting_end_datetime_local: str,
                      period_type: str,
                      language: str) -> Optional[str]:
        _ = self._

        if "reporting_period" not in report.keys() or \
                "names" not in report['reporting_period'].keys() or \
                len(report['reporting_period']['names']) == 0:
            doc = Document()
            section = doc.sections[0]
            section.orientation = 1
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            self._add_cover_page(doc, name, period_type,
                                 reporting_start_datetime_local,
                                 reporting_end_datetime_local,
                                 base_period_start_datetime_local,
                                 base_period_end_datetime_local,
                                 False)
            filename = str(uuid.uuid4()) + '.docx'
            doc.save(filename)
            return filename

        filename = str(uuid.uuid4()) + '.docx'

        self.report = _convert_decimals(report)
        self.name = name
        self.base_period_start = base_period_start_datetime_local
        self.base_period_end = base_period_end_datetime_local
        self.reporting_start = reporting_start_datetime_local
        self.reporting_end = reporting_end_datetime_local
        self.period_type = period_type

        self.is_base_period_exists = self._is_base_period_timestamp_exists(report['base_period'])

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        self._add_cover_page(doc, name, period_type,
                             reporting_start_datetime_local,
                             reporting_end_datetime_local,
                             base_period_start_datetime_local,
                             base_period_end_datetime_local,
                             self.is_base_period_exists)

        self._add_combined_analysis(doc)
        self._add_detailed_data_charts(doc)
        self._add_parameters_section(doc)

        doc.save(filename)
        logger.info(f"DOCX generated: {filename}")
        return filename

    def _add_heading_styled(self, doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        return heading

    def _add_cover_page(self, doc, name, period_type,
                        reporting_start, reporting_end,
                        base_period_start, base_period_end,
                        has_base_period):
        _ = self._
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                '..', 'excelexporters', 'myems.png')
        if os.path.exists(img_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(7.0))
            except Exception as e:
                logger.warning(f"Failed to load logo image: {e}")

        for _unused in range(3):
            doc.add_paragraph('')

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(_('Energy Plan Analysis'))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for _unused in range(3):
            doc.add_paragraph('')

        info_data = [
            [_('Name') + ':', name],
            [_('Period Type') + ':', period_type],
            [_('Reporting Start Datetime') + ':', reporting_start],
            [_('Reporting End Datetime') + ':', reporting_end],
        ]
        if has_base_period:
            info_data.append([_('Base Period Start Datetime') + ':', base_period_start])
            info_data.append([_('Base Period End Datetime') + ':', base_period_end])

        info_table = doc.add_table(rows=len(info_data), cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(info_table)

        half_w = Inches(3.2)
        for i, (label, value) in enumerate(info_data):
            c_label = info_table.cell(i, 0)
            c_label.width = half_w
            c_label.text = label
            c_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = c_label.paragraphs[0].runs[0]
            r_lbl.bold = True
            r_lbl.font.size = Pt(13)
            r_lbl.font.name = 'Arial'
            r_lbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            c_value = info_table.cell(i, 1)
            c_value.width = half_w
            c_value.text = str(value) if value is not None else ''
            c_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_val = c_value.paragraphs[0].runs[0]
            r_val.font.size = Pt(13)
            r_val.font.name = 'Arial'
            r_val._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        doc.add_page_break()

    def _add_combined_analysis(self, doc):
        """Add combined analysis: Reporting Period Plan summary table on top,
        TCE/TCO2E breakdown tables with pie charts below (two columns)."""
        _ = self._
        reporting_data = self.report['reporting_period']
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        subtotals_saving = reporting_data.get('subtotals_saving', [])
        subtotals_per_unit_area_saving = reporting_data.get('subtotals_per_unit_area_saving', [])
        increment_rates_saving = reporting_data.get('increment_rates_saving', [])
        subtotals_in_kgce_saving = reporting_data.get('subtotals_in_kgce_saving', [])
        subtotals_in_kgco2e_saving = reporting_data.get('subtotals_in_kgco2e_saving', [])
        ca_len = len(names)

        if ca_len == 0:
            return

        self._add_heading_styled(doc, self.name + ' - ' + _('Reporting Period Plan'), level=1)

        num_cols = ca_len + 3
        table = doc.add_table(rows=4, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['']
        for i in range(ca_len):
            unit_i = units[i] if (units and i < len(units)) else ''
            headers.append(names[i] + ' (' + _('Plan') + ' - ' + _('Actual') + ')(' + unit_i + ')')
        headers.append(_('Ton of Standard Coal') + ' (' + _('Plan') + ' - ' + _('Actual') + ') (TCE)')
        headers.append(_('Ton of Carbon Dioxide Emissions') + ' (' + _('Plan') + ' - ' + _('Actual') + ') (TCO2E)')

        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
            _style_table_cell(cell, is_header=True, bold=True, font_size=8)

        row_labels = [_('Plan'), _('Per Unit Area'), _('Increment Rate')]
        for r_idx, row_label in enumerate(row_labels, start=1):
            cell = table.cell(r_idx, 0)
            cell.text = row_label
            _style_table_cell(cell, is_green=True, bold=True, font_size=8)

        total_kgce = reporting_data.get('total_in_kgce_saving', 0)
        total_kgco2e = reporting_data.get('total_in_kgco2e_saving', 0)
        total_kgce_per_area = reporting_data.get('total_in_kgce_per_unit_area_saving', None)
        total_kgco2e_per_area = reporting_data.get('total_in_kgco2e_per_unit_area_saving', None)
        inc_kgce = reporting_data.get('increment_rate_in_kgce_saving', None)
        inc_kgco2e = reporting_data.get('increment_rate_in_kgco2e_saving', None)

        for i in range(ca_len):
            col = i + 1
            cell_plan = table.cell(1, col)
            val = subtotals_saving[i] if (subtotals_saving and i < len(subtotals_saving)) else None
            cell_plan.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(cell_plan, font_size=8)

            cell_area = table.cell(2, col)
            val = subtotals_per_unit_area_saving[i] if (subtotals_per_unit_area_saving and i < len(subtotals_per_unit_area_saving)) else None
            cell_area.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(cell_area, font_size=8)

            cell_inc = table.cell(3, col)
            val = increment_rates_saving[i] if (increment_rates_saving and i < len(increment_rates_saving)) else None
            cell_inc.text = (str(round2(val * 100, 2)) + '%') if val is not None else ''
            _style_table_cell(cell_inc, font_size=8)

        tce_col = ca_len + 1
        tco2e_col = ca_len + 2

        table.cell(1, tce_col).text = str(round2(total_kgce / 1000, 2))
        _style_table_cell(table.cell(1, tce_col), font_size=8)
        table.cell(1, tco2e_col).text = str(round2(total_kgco2e / 1000, 2))
        _style_table_cell(table.cell(1, tco2e_col), font_size=8)

        tce_area_text = (str(round2(total_kgce_per_area / 1000, 2))
                         if total_kgce_per_area is not None else '')
        table.cell(2, tce_col).text = tce_area_text
        _style_table_cell(table.cell(2, tce_col), font_size=8)
        tco2e_area_text = (str(round2(total_kgco2e_per_area / 1000, 2))
                           if total_kgco2e_per_area is not None else '')
        table.cell(2, tco2e_col).text = tco2e_area_text
        _style_table_cell(table.cell(2, tco2e_col), font_size=8)

        inc_tce_text = (str(round2(inc_kgce * 100, 2)) + '%') if inc_kgce is not None else ''
        table.cell(3, tce_col).text = inc_tce_text
        _style_table_cell(table.cell(3, tce_col), font_size=8)
        inc_tco2e_text = (str(round2(inc_kgco2e * 100, 2)) + '%') if inc_kgco2e is not None else ''
        table.cell(3, tco2e_col).text = inc_tco2e_text
        _style_table_cell(table.cell(3, tco2e_col), font_size=8)

        doc.add_paragraph('')

        self._add_tce_breakdown(doc, names, subtotals_in_kgce_saving, _('Plan'))
        self._add_tco2e_breakdown(doc, names, subtotals_in_kgco2e_saving, _('Plan'))

    def _add_tce_breakdown(self, doc, names, subtotals_in_kgce_saving, row_label_prefix):
        """Add TCE breakdown table + pie chart side-by-side."""
        _ = self._
        if not subtotals_in_kgce_saving or len(subtotals_in_kgce_saving) == 0:
            return
        ca_len = min(len(names), len(subtotals_in_kgce_saving))
        if ca_len == 0:
            return

        kgce_sum = sum(v or 0 for v in subtotals_in_kgce_saving)
        tce_values = [round2(subtotals_in_kgce_saving[i] / 1000, 3) for i in range(ca_len)]
        display_names = names[:ca_len]

        chart_file = self._make_pie_chart(tce_values, display_names,
                                          _('Ton of Standard Coal(TCE) by Energy Category'))

        container = doc.add_table(rows=1, cols=2)
        container.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(container)
        left_cell = container.cell(0, 0)
        right_cell = container.cell(0, 1)

        data_table = left_cell.add_table(rows=ca_len + 1, cols=4)
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        h0 = data_table.cell(0, 0)
        h0.text = ''
        _style_table_cell(h0, is_header=True, bold=True, font_size=8)
        h1 = data_table.cell(0, 1)
        h1.text = _('Energy Category')
        _style_table_cell(h1, is_header=True, bold=True, font_size=8)
        h2 = data_table.cell(0, 2)
        h2.text = row_label_prefix + ' (TCE)'
        _style_table_cell(h2, is_header=True, bold=True, font_size=8)
        h3 = data_table.cell(0, 3)
        h3.text = _('Ton of Standard Coal(TCE) by Energy Category')
        _style_table_cell(h3, is_header=True, bold=True, font_size=8)
        for i in range(ca_len):
            c0 = data_table.cell(i + 1, 0)
            c0.text = str(i + 1)
            _style_table_cell(c0, font_size=8)
            c1 = data_table.cell(i + 1, 1)
            c1.text = display_names[i]
            _style_table_cell(c1, bold=True, font_size=8)
            c2 = data_table.cell(i + 1, 2)
            c2.text = str(tce_values[i])
            _style_table_cell(c2, font_size=8)
            c3 = data_table.cell(i + 1, 3)
            proportion = (str(round2(subtotals_in_kgce_saving[i] / kgce_sum * 100, 2)) + '%'
                          if abs(kgce_sum) > 0 else '-')
            c3.text = proportion
            _style_table_cell(c3, font_size=8)

        if chart_file:
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_file, width=Inches(2.4))

    def _add_tco2e_breakdown(self, doc, names, subtotals_in_kgco2e_saving, row_label_prefix):
        """Add TCO2E breakdown table + pie chart side-by-side."""
        _ = self._
        if not subtotals_in_kgco2e_saving or len(subtotals_in_kgco2e_saving) == 0:
            return
        ca_len = min(len(names), len(subtotals_in_kgco2e_saving))
        if ca_len == 0:
            return

        kgco2e_sum = sum(v or 0 for v in subtotals_in_kgco2e_saving)
        co2e_values = [round2(subtotals_in_kgco2e_saving[i] / 1000, 3) for i in range(ca_len)]
        display_names = names[:ca_len]

        chart_file = self._make_pie_chart(co2e_values, display_names,
                                          _('Ton of Carbon Dioxide Emissions(TCO2E) by Energy Category'))

        container = doc.add_table(rows=1, cols=2)
        container.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(container)
        left_cell = container.cell(0, 0)
        right_cell = container.cell(0, 1)

        data_table = left_cell.add_table(rows=ca_len + 1, cols=4)
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        h0 = data_table.cell(0, 0)
        h0.text = ''
        _style_table_cell(h0, is_header=True, bold=True, font_size=8)
        h1 = data_table.cell(0, 1)
        h1.text = _('Energy Category')
        _style_table_cell(h1, is_header=True, bold=True, font_size=8)
        h2 = data_table.cell(0, 2)
        h2.text = row_label_prefix + ' (TCO2E)'
        _style_table_cell(h2, is_header=True, bold=True, font_size=8)
        h3 = data_table.cell(0, 3)
        h3.text = _('Ton of Carbon Dioxide Emissions(TCO2E) by Energy Category')
        _style_table_cell(h3, is_header=True, bold=True, font_size=8)
        for i in range(ca_len):
            c0 = data_table.cell(i + 1, 0)
            c0.text = str(i + 1)
            _style_table_cell(c0, font_size=8)
            c1 = data_table.cell(i + 1, 1)
            c1.text = display_names[i]
            _style_table_cell(c1, bold=True, font_size=8)
            c2 = data_table.cell(i + 1, 2)
            c2.text = str(co2e_values[i])
            _style_table_cell(c2, font_size=8)
            c3 = data_table.cell(i + 1, 3)
            proportion = (str(round2(subtotals_in_kgco2e_saving[i] / kgco2e_sum * 100, 2)) + '%'
                          if abs(kgco2e_sum) > 0 else '-')
            c3.text = proportion
            _style_table_cell(c3, font_size=8)

        if chart_file:
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_file, width=Inches(2.4))

    @staticmethod
    def _filter_valid_data(data):
        xs, ys = [], []
        for idx, v in enumerate(data):
            if v is not None:
                try:
                    ys.append(float(v))
                    xs.append(idx)
                except (TypeError, ValueError):
                    pass
        return xs, ys

    def _add_detailed_data_charts(self, doc):
        """Add detailed data line charts, up to 4 per page in 2x2 grid.
        Without base period: single line chart per category.
        With base period: solid/dashed comparison chart per category.
        """
        _ = self._

        reporting_data = self.report['reporting_period']
        timestamps = reporting_data.get('timestamps', [])
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        values_saving = reporting_data.get('values_saving', [])

        if not timestamps or len(timestamps[0]) == 0 or not names:
            return

        reporting_times = timestamps[0]
        num_categories = len(names)
        charts_per_page = 4

        doc.add_page_break()
        self._add_heading_styled(doc, self.name + ' ' + _('Detailed Data'), level=1)

        if not self.is_base_period_exists:
            all_charts = []
            for i in range(num_categories):
                raw_data = values_saving[i] if i < len(values_saving) else []
                xs, ys = self._filter_valid_data(raw_data)
                color = self.chart_colors[i % len(self.chart_colors)]

                fig, ax = plt.subplots(figsize=(4.8, 3.0))
                if ys:
                    ax.plot(xs, ys, linewidth=1.2, color=color,
                            marker='o', markersize=3,
                            markevery=max(1, len(ys) // 30))

                step = max(1, len(raw_data) // 10)
                ax.set_xticks(range(0, len(raw_data), step))
                ax.set_xticklabels(
                    [reporting_times[t][:10] for t in range(0, len(raw_data), step)],
                    rotation=45, ha='right', fontsize=7)
                unit_i = units[i] if (units and i < len(units)) else ''
                ax.set_title(_('Reporting Period Plan') + ' - ' +
                             names[i] + ' (' + unit_i + ')',
                             fontsize=9, fontweight='bold')
                ax.grid(True, alpha=0.3)
                all_charts.append(self._fig_to_bytesio(fig, self.dpi))

            num_total_charts = len(all_charts)
            for page_start in range(0, num_total_charts, charts_per_page):
                page_end = min(page_start + charts_per_page, num_total_charts)
                page_indices = list(range(page_start, page_end))
                num_on_page = len(page_indices)

                rows = (num_on_page + 1) // 2

                for row_idx in range(rows):
                    slot0 = row_idx * 2
                    slot1 = slot0 + 1
                    has_left = slot0 < num_on_page
                    has_right = slot1 < num_on_page

                    if has_left and not has_right:
                        container = doc.add_table(rows=1, cols=2)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)
                        container.cell(0, 0).merge(container.cell(0, 1))
                        cell = container.cell(0, 0)
                        chart_buf = all_charts[page_indices[slot0]]
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(chart_buf, width=Inches(4.8))
                    else:
                        container_cols = min(2, num_on_page - row_idx * 2)
                        container = doc.add_table(rows=1, cols=container_cols)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)

                        for col_idx in range(container_cols):
                            slot = row_idx * 2 + col_idx
                            if slot >= num_on_page:
                                continue
                            chart_buf = all_charts[page_indices[slot]]

                            cell = container.cell(0, col_idx)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(chart_buf, width=Inches(4.8))

                if page_end < num_total_charts:
                    doc.add_page_break()
        else:
            base_period_data = self.report['base_period']
            base_values_saving = base_period_data.get('values_saving', [])

            all_charts = []
            for i in range(num_categories):
                color = self.chart_colors[i % len(self.chart_colors)]

                fig, ax = plt.subplots(figsize=(4.8, 3.0))

                r_data = values_saving[i] if i < len(values_saving) else []
                r_xs, r_ys = self._filter_valid_data(r_data)
                if r_ys:
                    ax.plot(r_xs, r_ys, linewidth=1.2, color=color,
                            marker='o', markersize=3,
                            markevery=max(1, len(r_ys) // 30),
                            label=_('Reporting Period') + ' - ' + names[i])

                if i < len(base_values_saving):
                    b_data = base_values_saving[i]
                    b_xs, b_ys = self._filter_valid_data(b_data)
                    if b_ys:
                        ax.plot(b_xs, b_ys, linewidth=1.2, color=color,
                                linestyle='--', marker='s', markersize=3,
                                markevery=max(1, len(b_ys) // 30),
                                label=_('Base Period') + ' - ' + names[i])

                step = max(1, len(r_data) // 10)
                ax.set_xticks(range(0, len(r_data), step))
                ax.set_xticklabels(
                    [reporting_times[t][:10] if t < len(reporting_times) else ''
                     for t in range(0, len(r_data), step)],
                    rotation=45, ha='right', fontsize=7)
                unit_i = units[i] if (units and i < len(units)) else ''
                ax.set_title(
                    _('Base Period Plan') + ' / ' +
                    _('Reporting Period Plan') + ' - ' +
                    names[i] + ' (' + unit_i + ')',
                    fontsize=8, fontweight='bold')
                if r_ys or (i < len(base_values_saving) and b_ys):
                    ax.legend(fontsize=7)
                ax.grid(True, alpha=0.3)
                all_charts.append(self._fig_to_bytesio(fig, self.dpi))

            num_total_charts = len(all_charts)
            for page_start in range(0, num_total_charts, charts_per_page):
                page_end = min(page_start + charts_per_page, num_total_charts)
                page_indices = list(range(page_start, page_end))
                num_on_page = len(page_indices)

                rows = (num_on_page + 1) // 2

                for row_idx in range(rows):
                    slot0 = row_idx * 2
                    slot1 = slot0 + 1
                    has_left = slot0 < num_on_page
                    has_right = slot1 < num_on_page

                    if has_left and not has_right:
                        container = doc.add_table(rows=1, cols=2)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)
                        container.cell(0, 0).merge(container.cell(0, 1))
                        cell = container.cell(0, 0)
                        chart_buf = all_charts[page_indices[slot0]]
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(chart_buf, width=Inches(4.8))
                    else:
                        container_cols = min(2, num_on_page - row_idx * 2)
                        container = doc.add_table(rows=1, cols=container_cols)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)

                        for col_idx in range(container_cols):
                            slot = row_idx * 2 + col_idx
                            if slot >= num_on_page:
                                continue
                            chart_buf = all_charts[page_indices[slot]]

                            cell = container.cell(0, col_idx)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(chart_buf, width=Inches(4.8))

                if page_end < num_total_charts:
                    doc.add_page_break()

        doc.add_page_break()

    def _add_parameters_section(self, doc):
        """Add parameter data tables and filled line charts."""
        _ = self._

        params = self.report.get('parameters', {})
        if not params or not params.get('names') or not params.get('timestamps'):
            return

        param_names = params.get('names', [])
        timestamps = params.get('timestamps', [])
        values = params.get('values', [])
        units = params.get('units', [])

        all_zero = True
        for ts_list in timestamps:
            if ts_list and len(ts_list) > 0:
                all_zero = False
                break
        if all_zero:
            return

        all_params = list(range(len(param_names)))
        valid_params = []
        for i in all_params:
            if i < len(timestamps) and len(timestamps[i]) > 0:
                if i < len(values) and len(values[i]) > 0:
                    valid_params.append(i)
        if not valid_params:
            return

        self._add_heading_styled(doc, self.name + ' ' + _('Parameters'), level=1)

        rows_per_param = 10

        for pi in valid_params:
            name = param_names[pi]
            unit_i = units[pi] if (units and pi < len(units)) else ''
            times = timestamps[pi]
            data = values[pi]
            data_len = len(times)

            display_name = name + ((' (' + unit_i + ')') if unit_i else '')

            tbl_rows = min(rows_per_param, data_len)

            fig, ax = plt.subplots(figsize=(5.0, 2.4))
            marker_step_p = max(1, data_len // 20)
            color = '#5B9BD5'
            ax.plot(range(data_len), data, linewidth=1.2,
                    color=color, marker='o', markersize=3,
                    markevery=marker_step_p, label=name)
            ax.fill_between(range(data_len), data, alpha=0.15, color=color)
            step = max(1, data_len // 8)
            ax.set_xticks(range(0, data_len, step))
            ax.set_xticklabels([times[t][:10] for t in range(0, data_len, step)],
                               rotation=45, ha='right', fontsize=6)
            ax.set_ylabel(name, fontsize=8)
            ax.set_title(display_name, fontsize=9, fontweight='bold')
            ax.grid(True, alpha=0.3)
            chart_buf = self._fig_to_bytesio(fig, self.dpi)

            container = doc.add_table(rows=1, cols=2)
            container.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(container)
            left_cell = container.cell(0, 0)
            right_cell = container.cell(0, 1)

            data_table = left_cell.add_table(rows=tbl_rows + 1, cols=2)
            data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            h0 = data_table.cell(0, 0)
            h0.text = _('Time')
            _style_table_cell(h0, is_header=True, bold=True, font_size=8)
            h1 = data_table.cell(0, 1)
            h1.text = name
            _style_table_cell(h1, is_header=True, bold=True, font_size=8)
            for j in range(tbl_rows):
                c_t = data_table.cell(j + 1, 0)
                c_t.text = str(times[j])
                _style_table_cell(c_t, font_size=7)
                c_v = data_table.cell(j + 1, 1)
                c_v.text = str(round2(data[j], 2))
                _style_table_cell(c_v, font_size=7)

            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_buf, width=Inches(4.5))

    def _is_base_period_timestamp_exists(self, base_period_data: Dict) -> bool:
        timestamps = base_period_data.get('timestamps', [])

        if not timestamps:
            return False

        for timestamp in timestamps:
            if timestamp and len(timestamp) > 0:
                return True

        return False


def export(report,
           name,
           base_period_start_datetime_local,
           base_period_end_datetime_local,
           reporting_start_datetime_local,
           reporting_end_datetime_local,
           period_type,
           language):
    """
    Export report data to DOCX and return base64 encoded string.
    This function maintains the same interface as the Excel exporter.
    """
    exporter = SpacePlanDOCXExporter(language)
    return exporter.export(report, name,
                           base_period_start_datetime_local,
                           base_period_end_datetime_local,
                           reporting_start_datetime_local,
                           reporting_end_datetime_local,
                           period_type,
                           language)
