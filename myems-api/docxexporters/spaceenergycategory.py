"""
Space Energy Category DOCX Exporter

This module provides functionality to export space energy category data to DOCX format.
It generates comprehensive reports showing energy consumption breakdown by categories
for spaces with detailed analysis and visualizations.

Key Features:
- Space energy consumption by category
- Base period vs reporting period comparison
- Energy category proportion analysis
- Detailed data with charts
- Child spaces data with charts
- Parameter data (if available)
- Multi-language support
- Base64 encoding for file transmission

The exported DOCX file includes:
- Cover page with logo and report metadata
- Reporting period consumption summary (values, per unit area, increment rate)
- Time-of-use electricity consumption table + pie chart
- TCE (Ton of Standard Coal) breakdown with pie chart
- TCO2E (Ton of CO2 Equivalent) breakdown with pie chart
- Child spaces data table + per-category pie charts
- Base period working / non-working days comparison table
- Reporting period working / non-working days comparison table
- Detailed time-series paginated tables + per-page trend line charts
  (two layouts: without base period, with base period side-by-side comparison)
- Parameter data (sensors, tariffs, etc.) tables + line charts with filled area
"""

import base64
import os
import time
import uuid
import io

from decimal import Decimal
from typing import Optional, Dict, List, Any, BinaryIO
import logging

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.utilities import get_translation, round2

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

_font_setup_done = False


def setup_chinese_fonts():
    global _font_setup_done
    if _font_setup_done:
        return True

    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             '..', 'pdfexporters', 'fonts', 'NotoSansCJK-Regular.ttc')
    try:
        import matplotlib.font_manager as fm
        fm.fontManager.addfont(font_path)
        prop = fm.FontProperties(fname=font_path)
        font_name = prop.get_name()
        plt.rcParams['font.sans-serif'] = [font_name]
        plt.rcParams['axes.unicode_minus'] = False
        logger.info(f"Successfully loaded bundled font: {font_name} from {font_path}")
        _font_setup_done = True
        return True
    except Exception as e:
        logger.warning(f"Failed to load bundled font from {font_path}: {e}")

    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    logger.warning("Failed to load bundled NotoSansCJK font, using DejaVu Sans")
    return False


def _convert_decimals(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    elif isinstance(obj, dict):
        return {k: _convert_decimals(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_convert_decimals(item) for item in obj]
    elif isinstance(obj, tuple):
        return tuple(_convert_decimals(item) for item in obj)
    return obj


def _set_min_row_height(cell, height_twips=150, exact=False):
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def _reset_cell_paragraph_spacing(cell, font_size=9):
    line_twips = max(int(font_size * 20 * 1.15), 120)
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(spacing)
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


def _set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'start', 'bottom', 'end']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _style_table_cell(cell, is_header=False, is_green=False, bold=False, font_size=9):
    cell.paragraphs[0].runs[0].font.bold = bold
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _reset_cell_paragraph_spacing(cell, font_size=font_size)
    _set_cell_margins_zero(cell)
    row_h = max(int(font_size * 20 * 1.3), 160)
    _set_min_row_height(cell, height_twips=row_h, exact=False)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Borders
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    # Background color
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    if is_header:
        shd.set(qn('w:fill'), 'D9E2F3')
    elif is_green:
        shd.set(qn('w:fill'), '90EE90')
    else:
        return
    tcPr.append(shd)


class SpaceEnergyDOCXExporter:
    """
    Export space energy category data to DOCX format.
    Generates comprehensive reports with charts and tables matching Excel layout.
    """

    def __init__(self, language: str = 'zh_CN'):
        font_setup_success = setup_chinese_fonts()
        if not font_setup_success:
            logger.warning("Chinese font setup failed, some text may not display correctly")

        self.language = language
        self.trans = get_translation(language)
        self._ = self.trans.gettext

        self.dpi = 120

        self.chart_colors = ['#4472C4', '#ED7D31', '#70AD47', '#FFC000', '#5B9BD5',
                             '#FF6B6B', '#9B59B6', '#1ABC9C', '#E67E22', '#2ECC71',
                             '#3498DB', '#E74C3C', '#2ECC71', '#F39C12', '#9B59B6']

    def export(self,
               report: Dict[str, Any],
               name: str,
               base_period_start_datetime_local: str,
               base_period_end_datetime_local: str,
               reporting_start_datetime_local: str,
               reporting_end_datetime_local: str,
               period_type: str,
               language: str) -> Optional[str]:
        if report is None:
            return None
        start_time = time.time()
        logger.info(f"Starting DOCX generation for {name}")

        docx_filename = self.generate_docx(
            report, name,
            base_period_start_datetime_local,
            base_period_end_datetime_local,
            reporting_start_datetime_local,
            reporting_end_datetime_local,
            period_type,
            language
        )

        result = ''
        if docx_filename and os.path.exists(docx_filename):
            try:
                with open(docx_filename, 'rb') as binary_file:
                    binary_data = binary_file.read()
                result = base64.b64encode(binary_data).decode('utf-8')
            except Exception as e:
                logger.error(f"Failed to encode DOCX: {str(e)}")
            finally:
                try:
                    os.remove(docx_filename)
                except Exception:
                    pass
        elapsed = time.time() - start_time
        logger.info(f"DOCX generation completed in {elapsed:.2f}s for {name}")
        return result

    @staticmethod
    def _fig_to_bytesio(fig, dpi: int) -> BinaryIO:
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    def _make_pie_chart(self, values, labels, title, colors=None):
        if not values or sum((v or 0) for v in values) == 0:
            return None
        fig, ax = plt.subplots(figsize=(3.2, 2.6))
        if colors is None:
            colors = self.chart_colors[:len(labels)]
        filtered = [(l, v, c) for l, v, c in zip(labels, values, colors) if (v or 0) > 0]
        if not filtered:
            plt.close(fig)
            return None
        f_labels, f_values, f_colors = zip(*filtered)
        if len(f_labels) > 8:
            sorted_data = sorted(zip(f_labels, f_values, f_colors), key=lambda x: x[1], reverse=True)
            top_data = sorted_data[:7]
            other_sum = sum(v for _, v, _ in sorted_data[7:])
            f_labels = [l for l, _, _ in top_data] + [self._('Others')]
            f_values = [v for _, v, _ in top_data] + [other_sum]
            f_colors = [c for _, _, c in top_data] + ['#999999']
        ax.pie(f_values, labels=f_labels, autopct='%1.1f%%', colors=f_colors, startangle=90)
        ax.set_title(title, fontsize=10, fontweight='bold')
        return self._fig_to_bytesio(fig, self.dpi)

    def _make_pie_charts_row(self, charts):
        if not charts:
            return None
        n = len(charts)
        per_w = 3.2
        fig, axes = plt.subplots(1, n, figsize=(per_w * n, 3.2))
        if n == 1:
            axes = [axes]
        plotted_any = False
        for idx, chart in enumerate(charts):
            ax = axes[idx]
            values = chart['values']
            labels = chart['labels']
            title = chart['title']
            colors = chart.get('colors')
            if not values or sum((v or 0) for v in values) == 0:
                ax.axis('off')
                continue
            if colors is None:
                colors = self.chart_colors[:len(labels)]
            filtered = [(l, v, c) for l, v, c in zip(labels, values, colors) if (v or 0) > 0]
            if not filtered:
                ax.axis('off')
                continue
            f_labels, f_values, f_colors = zip(*filtered)
            if len(f_labels) > 8:
                sorted_data = sorted(zip(f_labels, f_values, f_colors), key=lambda x: x[1], reverse=True)
                top_data = sorted_data[:7]
                other_sum = sum(v for _, v, _ in sorted_data[7:])
                f_labels = [l for l, _, _ in top_data] + [self._('Others')]
                f_values = [v for _, v, _ in top_data] + [other_sum]
                f_colors = [c for _, _, c in top_data] + ['#999999']
            ax.pie(f_values, labels=f_labels, autopct='%1.1f%%', colors=f_colors, startangle=90)
            ax.set_title(title, fontsize=10, fontweight='bold')
            plotted_any = True
        if not plotted_any:
            plt.close(fig)
            return None
        plt.tight_layout(pad=1.0)
        return self._fig_to_bytesio(fig, self.dpi)

    def generate_docx(self,
                      report: Dict[str, Any],
                      name: str,
                      base_period_start_datetime_local: str,
                      base_period_end_datetime_local: str,
                      reporting_start_datetime_local: str,
                      reporting_end_datetime_local: str,
                      period_type: str,
                      language: str) -> Optional[str]:
        _ = self._

        # Check if there is data; even empty still generate with cover
        if "reporting_period" not in report.keys() or \
                "names" not in report['reporting_period'].keys() or \
                len(report['reporting_period']['names']) == 0:
            doc = Document()
            section = doc.sections[0]
            section.orientation = 1  # landscape
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            self._add_cover_page(doc, name, period_type,
                                 reporting_start_datetime_local,
                                 reporting_end_datetime_local,
                                 base_period_start_datetime_local,
                                 base_period_end_datetime_local,
                                 False)
            filename = str(uuid.uuid4()) + '.docx'
            doc.save(filename)
            return filename

        filename = str(uuid.uuid4()) + '.docx'

        # Prepare data
        self.report = _convert_decimals(report)
        self.name = name
        self.base_period_start = base_period_start_datetime_local
        self.base_period_end = base_period_end_datetime_local
        self.reporting_start = reporting_start_datetime_local
        self.reporting_end = reporting_end_datetime_local
        self.period_type = period_type

        self.is_base_period_exists = self._is_base_period_timestamp_exists(report['base_period'])

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        self._add_cover_page(doc, name, period_type,
                             reporting_start_datetime_local,
                             reporting_end_datetime_local,
                             base_period_start_datetime_local,
                             base_period_end_datetime_local,
                             self.is_base_period_exists)

        self._add_reporting_period_summary(doc)
        self._add_time_of_use_section(doc)
        self._add_tce_section(doc)
        self._add_tco2e_section(doc)
        self._add_child_spaces_section(doc)
        self._add_base_period_working_days_section(doc)
        self._add_reporting_working_days_section(doc)
        self._add_detailed_data_section(doc)
        self._add_parameters_section(doc)

        doc.save(filename)
        logger.info(f"DOCX generated: {filename}")
        return filename

    # ---------- Utility ----------
    def _add_heading_styled(self, doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        return heading

    def _add_cover_page(self, doc, name, period_type,
                        reporting_start, reporting_end,
                        base_period_start, base_period_end,
                        has_base_period):
        _ = self._
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                '..', 'excelexporters', 'myems.png')
        if os.path.exists(img_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(3.5))
            except Exception as e:
                logger.warning(f"Failed to load logo image: {e}")

        for _unused in range(3):
            doc.add_paragraph('')

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(_('Space Data - Energy Category Analysis'))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for _unused in range(3):
            doc.add_paragraph('')

        info_data = [
            [_('Name') + ':', name],
            [_('Period Type') + ':', period_type],
            [_('Reporting Start Datetime') + ':', reporting_start],
            [_('Reporting End Datetime') + ':', reporting_end],
        ]
        if has_base_period:
            info_data.append([_('Base Period Start Datetime') + ':', base_period_start])
            info_data.append([_('Base Period End Datetime') + ':', base_period_end])

        table = doc.add_table(rows=len(info_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(info_data):
            cell_label = table.cell(i, 0)
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].font.bold = True
            cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            cell_value = table.cell(i, 1)
            cell_value.text = str(value) if value is not None else ''
            cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ---------- Combined analysis equivalent ----------
    def _add_reporting_period_summary(self, doc):
        """Add reporting period consumption summary table matching Excel layout."""
        _ = self._
        reporting_data = self.report['reporting_period']
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        subtotals = reporting_data.get('subtotals', [])
        subtotals_per_unit_area = reporting_data.get('subtotals_per_unit_area', [])
        increment_rates = reporting_data.get('increment_rates', [])
        ca_len = len(names)

        if ca_len == 0:
            return

        self._add_heading_styled(doc, self.name + ' - ' + _('Reporting Period Consumption'), level=1)

        num_cols = ca_len + 3
        table = doc.add_table(rows=4, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['']
        for i in range(ca_len):
            unit_i = units[i] if (units and i < len(units)) else ''
            headers.append(names[i] + ((' (' + unit_i + ')') if unit_i else ''))
        headers.append(_('Ton of Standard Coal') + '(TCE)')
        headers.append(_('Ton of Carbon Dioxide Emissions') + '(TCO2E)')

        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
            _style_table_cell(cell, is_header=True, bold=True)

        row_labels = [_('Consumption'), _('Per Unit Area'), _('Increment Rate')]
        for r_idx, row_label in enumerate(row_labels, start=1):
            cell = table.cell(r_idx, 0)
            cell.text = row_label
            _style_table_cell(cell, is_green=True, bold=True)

        total_kgce = reporting_data.get('total_in_kgce', 0)
        total_kgco2e = reporting_data.get('total_in_kgco2e', 0)
        total_kgce_per_area = reporting_data.get('total_in_kgce_per_unit_area', None)
        total_kgco2e_per_area = reporting_data.get('total_in_kgco2e_per_unit_area', None)
        inc_kgce = reporting_data.get('increment_rate_in_kgce', None)
        inc_kgco2e = reporting_data.get('increment_rate_in_kgco2e', None)

        for i in range(ca_len):
            col = i + 1
            cell_cons = table.cell(1, col)
            val = subtotals[i] if (subtotals and i < len(subtotals)) else None
            cell_cons.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(cell_cons)

            cell_area = table.cell(2, col)
            val = subtotals_per_unit_area[i] if (subtotals_per_unit_area and i < len(subtotals_per_unit_area)) else None
            cell_area.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(cell_area)

            cell_inc = table.cell(3, col)
            val = increment_rates[i] if (increment_rates and i < len(increment_rates)) else None
            cell_inc.text = (str(round2(val * 100, 2)) + '%') if val is not None else ''
            _style_table_cell(cell_inc)

        tce_col = ca_len + 1
        tco2e_col = ca_len + 2

        table.cell(1, tce_col).text = str(round2(total_kgce / 1000, 2))
        _style_table_cell(table.cell(1, tce_col))
        table.cell(1, tco2e_col).text = str(round2(total_kgco2e / 1000, 2))
        _style_table_cell(table.cell(1, tco2e_col))

        tce_area_text = (str(round2(total_kgce_per_area / 1000, 2))
                         if total_kgce_per_area is not None else '')
        table.cell(2, tce_col).text = tce_area_text
        _style_table_cell(table.cell(2, tce_col))
        tco2e_area_text = (str(round2(total_kgco2e_per_area / 1000, 2))
                           if total_kgco2e_per_area is not None else '')
        table.cell(2, tco2e_col).text = tco2e_area_text
        _style_table_cell(table.cell(2, tco2e_col))

        inc_tce_text = (str(round2(inc_kgce * 100, 2)) + '%') if inc_kgce is not None else ''
        table.cell(3, tce_col).text = inc_tce_text
        _style_table_cell(table.cell(3, tce_col))
        inc_tco2e_text = (str(round2(inc_kgco2e * 100, 2)) + '%') if inc_kgco2e is not None else ''
        table.cell(3, tco2e_col).text = inc_tco2e_text
        _style_table_cell(table.cell(3, tco2e_col))

    def _add_time_of_use_section(self, doc):
        """Add Time-Of-Use electricity consumption table and pie chart."""
        _ = self._
        reporting_data = self.report['reporting_period']
        electricity_index = -1
        for i in range(len(reporting_data.get('energy_category_ids', []))):
            if reporting_data['energy_category_ids'][i] == 1:
                electricity_index = i
                break
        if electricity_index < 0:
            return

        toppeaks = reporting_data.get('toppeaks', [])
        onpeaks = reporting_data.get('onpeaks', [])
        midpeaks = reporting_data.get('midpeaks', [])
        offpeaks = reporting_data.get('offpeaks', [])
        if not toppeaks or len(toppeaks) <= electricity_index:
            return

        tou_values = [
            round2(toppeaks[electricity_index], 2) if electricity_index < len(toppeaks) else 0,
            round2(onpeaks[electricity_index], 2) if electricity_index < len(onpeaks) else 0,
            round2(midpeaks[electricity_index], 2) if electricity_index < len(midpeaks) else 0,
            round2(offpeaks[electricity_index], 2) if electricity_index < len(offpeaks) else 0,
        ]
        tou_categories = [_('TopPeak'), _('OnPeak'), _('MidPeak'), _('OffPeak')]
        tou_colors = ['#FF1744', '#FF6F00', '#FDD835', '#00BCD4']

        chart_file = self._make_pie_chart(tou_values, tou_categories,
                                          _('Electricity Consumption by Time-Of-Use'), colors=tou_colors)

        container = doc.add_table(rows=1, cols=2)
        container.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(container)
        left_cell = container.cell(0, 0)
        right_cell = container.cell(0, 1)

        data_table = left_cell.add_table(rows=5, cols=2)
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        h1 = data_table.cell(0, 0)
        h1.text = ''
        _style_table_cell(h1, is_header=True, bold=True)
        h2 = data_table.cell(0, 1)
        h2.text = _('Electricity Consumption by Time-Of-Use')
        _style_table_cell(h2, is_header=True, bold=True)
        for i in range(4):
            c1 = data_table.cell(i + 1, 0)
            c1.text = tou_categories[i]
            _style_table_cell(c1, bold=True)
            c2 = data_table.cell(i + 1, 1)
            c2.text = str(tou_values[i])
            _style_table_cell(c2)

        if chart_file:
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_file, width=Inches(2.0))

    def _add_tce_section(self, doc):
        """Add Ton of Standard Coal(TCE) breakdown table and pie chart by energy category."""
        _ = self._
        reporting_data = self.report['reporting_period']
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        subtotals_in_kgce = reporting_data.get('subtotals_in_kgce', [])
        if not subtotals_in_kgce or len(subtotals_in_kgce) == 0:
            return
        if sum((v or 0) for v in subtotals_in_kgce) == 0:
            return

        ca_len = min(len(names), len(subtotals_in_kgce))
        if ca_len == 0:
            return
        tce_values = [round2(subtotals_in_kgce[i] / 1000, 3) for i in range(ca_len)]
        display_names = names[:ca_len]

        chart_file = self._make_pie_chart(tce_values, display_names,
                                          _('Ton of Standard Coal(TCE) by Energy Category'))

        container = doc.add_table(rows=1, cols=2)
        container.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(container)
        left_cell = container.cell(0, 0)
        right_cell = container.cell(0, 1)

        data_table = left_cell.add_table(rows=ca_len + 1, cols=2)
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        h1 = data_table.cell(0, 0)
        h1.text = ''
        _style_table_cell(h1, is_header=True, bold=True)
        h2 = data_table.cell(0, 1)
        h2.text = _('Ton of Standard Coal(TCE) by Energy Category')
        _style_table_cell(h2, is_header=True, bold=True)
        for i in range(ca_len):
            c1 = data_table.cell(i + 1, 0)
            c1.text = display_names[i]
            _style_table_cell(c1, bold=True)
            c2 = data_table.cell(i + 1, 1)
            c2.text = str(tce_values[i])
            _style_table_cell(c2)

        if chart_file:
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_file, width=Inches(2.0))

    def _add_tco2e_section(self, doc):
        """Add Ton of CO2 Equivalent(TCO2E) breakdown table and pie chart by energy category."""
        _ = self._
        reporting_data = self.report['reporting_period']
        names = reporting_data.get('names', [])
        subtotals_in_kgco2e = reporting_data.get('subtotals_in_kgco2e', [])
        if not subtotals_in_kgco2e or len(subtotals_in_kgco2e) == 0:
            return
        if sum((v or 0) for v in subtotals_in_kgco2e) == 0:
            return

        ca_len = min(len(names), len(subtotals_in_kgco2e))
        if ca_len == 0:
            return
        co2e_values = [round2(subtotals_in_kgco2e[i] / 1000, 3) for i in range(ca_len)]
        display_names = names[:ca_len]

        chart_file = self._make_pie_chart(co2e_values, display_names,
                                          _('Ton of Carbon Dioxide Emissions(TCO2E) by Energy Category'))

        container = doc.add_table(rows=1, cols=2)
        container.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(container)
        left_cell = container.cell(0, 0)
        right_cell = container.cell(0, 1)

        data_table = left_cell.add_table(rows=ca_len + 1, cols=2)
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        h1 = data_table.cell(0, 0)
        h1.text = ''
        _style_table_cell(h1, is_header=True, bold=True)
        h2 = data_table.cell(0, 1)
        h2.text = _('Ton of Carbon Dioxide Emissions(TCO2E) by Energy Category')
        _style_table_cell(h2, is_header=True, bold=True)
        for i in range(ca_len):
            c1 = data_table.cell(i + 1, 0)
            c1.text = display_names[i]
            _style_table_cell(c1, bold=True)
            c2 = data_table.cell(i + 1, 1)
            c2.text = str(co2e_values[i])
            _style_table_cell(c2)

        if chart_file:
            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_file, width=Inches(2.0))

    # ---------- Child spaces ----------
    def _add_child_spaces_section(self, doc):
        """Add child spaces data table with per-category pie charts matching Excel layout."""
        _ = self._

        child = self.report.get('child_space', {})
        if not child or 'energy_category_names' not in child or not child['energy_category_names']:
            return
        if 'child_space_ids_array' not in child or 'child_space_names_array' not in child:
            return
        if not child['child_space_ids_array'] or not child['child_space_names_array']:
            return
        if not child['child_space_names_array'][0] or len(child['child_space_names_array'][0]) == 0:
            return

        names_array = child['child_space_names_array']
        child_names = names_array[0]
        child_ids = child['child_space_ids_array'][0]
        category_names = child['energy_category_names']
        units = child.get('units', [])
        subtotals_array = child['subtotals_array']
        ca_len = len(category_names)
        space_len = len(child_names)

        self._add_heading_styled(doc, self.name + ' ' + _('Child Spaces Data'), level=1)

        num_cols = 2 + ca_len * 2
        table = doc.add_table(rows=space_len + 1, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        c0 = table.cell(0, 0)
        c0.text = _('ID')
        _style_table_cell(c0, is_header=True, bold=True)
        c1 = table.cell(0, 1)
        c1.text = _('Child Space')
        _style_table_cell(c1, is_header=True, bold=True)
        for j in range(ca_len):
            col = 2 + j * 2
            hdr1 = table.cell(0, col)
            unit_j = units[j] if (units and j < len(units)) else ''
            hdr1.text = category_names[j] + ((' (' + unit_j + ')') if unit_j else '')
            _style_table_cell(hdr1, is_header=True, bold=True)
            hdr2 = table.cell(0, col + 1)
            hdr2.text = ''
            _style_table_cell(hdr2, is_header=True, bold=True)

        for i in range(space_len):
            r_idx = i + 1
            c_id = table.cell(r_idx, 0)
            c_id.text = str(child_ids[i])
            _style_table_cell(c_id)
            c_name = table.cell(r_idx, 1)
            c_name.text = child_names[i]
            _style_table_cell(c_name)
            for j in range(ca_len):
                col = 2 + j * 2
                total = sum((v or 0) for v in subtotals_array[j]) if subtotals_array[j] else 0
                val = round2(subtotals_array[j][i], 2) if i < len(subtotals_array[j]) else 0
                c_val = table.cell(r_idx, col)
                c_val.text = str(val)
                _style_table_cell(c_val)
                pct = str(round2(val / total * 100, 2)) + '%' if total > 0 else '0.00%'
                c_pct = table.cell(r_idx, col + 1)
                c_pct.text = pct
                _style_table_cell(c_pct)

        charts_per_row = 3
        per_chart_w = 3.2
        charts_data = []
        for j in range(ca_len):
            values = []
            for s in range(space_len):
                v = subtotals_array[j][s] if s < len(subtotals_array[j]) else 0
                values.append(v or 0)
            labels = child_names[:len(values)]
            unit_j = units[j] if (units and j < len(units)) else ''
            chart_title = category_names[j] + ((' (' + unit_j + ')') if unit_j else '')
            charts_data.append({'values': values, 'labels': labels, 'title': chart_title})

        doc.add_paragraph('')

        num_rows = (ca_len + charts_per_row - 1) // charts_per_row
        for row_idx in range(num_rows):
            start = row_idx * charts_per_row
            end = min(start + charts_per_row, ca_len)
            row_charts = charts_data[start:end]
            row_file = self._make_pie_charts_row(row_charts)
            if not row_file:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            total_w = per_chart_w * len(row_charts)
            max_w = 10
            insert_w = min(total_w, max_w)
            run.add_picture(row_file, width=Inches(insert_w))

    # ---------- Working days (Base & Reporting) ----------
    def _add_base_period_working_days_section(self, doc):
        """Add base period working/non-working days comparison table matching Excel layout."""
        _ = self._

        base_period = self.report.get('base_period', {})
        if not self.is_base_period_exists:
            return

        non_working = base_period.get('non_working_days_subtotals', [])
        working = base_period.get('working_days_subtotals', [])
        names = base_period.get('names', [])
        units = base_period.get('units', [])

        if not working or not non_working:
            return
        if sum((v or 0) for v in working) == 0 and sum((v or 0) for v in non_working) == 0:
            return

        ca_len = len(names)

        self._add_heading_styled(doc, self.name + ' ' + _('Base Period Consumption'), level=1)

        col_headers = ['', _('Non Working Days') + _('Consumption'),
                       _('Working Days') + _('Consumption')]
        table = doc.add_table(rows=ca_len + 1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(col_headers):
            c = table.cell(0, j)
            c.text = h
            _style_table_cell(c, is_header=True, bold=True)

        space_working_calendars = self.report.get('space', {}).get('working_calendars', [])

        for i in range(ca_len):
            unit_i = units[i] if (units and i < len(units)) else ''
            label = names[i] + ((' (' + unit_i + ')') if unit_i else '')
            nw_val = non_working[i] if i < len(non_working) else 0
            w_val = working[i] if i < len(working) else 0
            nw_display = str(nw_val) if len(space_working_calendars) > 0 and (nw_val or 0) > 0 else '-'
            w_display = str(w_val) if len(space_working_calendars) > 0 and (w_val or 0) > 0 else '-'
            c0 = table.cell(i + 1, 0)
            c0.text = label
            _style_table_cell(c0, bold=True)
            c1 = table.cell(i + 1, 1)
            c1.text = nw_display
            _style_table_cell(c1)
            c2 = table.cell(i + 1, 2)
            c2.text = w_display
            _style_table_cell(c2)

    def _add_reporting_working_days_section(self, doc):
        """Add reporting period working/non-working days comparison table matching Excel layout."""
        _ = self._

        reporting_period = self.report.get('reporting_period', {})
        non_working = reporting_period.get('non_working_days_subtotals', [])
        working = reporting_period.get('working_days_subtotals', [])
        names = reporting_period.get('names', [])
        units = reporting_period.get('units', [])

        if not working or not non_working:
            return
        if sum((v or 0) for v in working) == 0 and sum((v or 0) for v in non_working) == 0:
            return

        ca_len = len(names)

        self._add_heading_styled(doc, self.name + ' ' + _('Reporting Period Consumption'), level=1)

        col_headers = ['', _('Non Working Days') + _('Consumption'),
                       _('Working Days') + _('Consumption')]
        table = doc.add_table(rows=ca_len + 1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(col_headers):
            c = table.cell(0, j)
            c.text = h
            _style_table_cell(c, is_header=True, bold=True)

        space_working_calendars = self.report.get('space', {}).get('working_calendars', [])

        for i in range(ca_len):
            unit_i = units[i] if (units and i < len(units)) else ''
            label = names[i] + ((' (' + unit_i + ')') if unit_i else '')
            nw_val = non_working[i] if i < len(non_working) else 0
            w_val = working[i] if i < len(working) else 0
            nw_display = str(nw_val) if len(space_working_calendars) > 0 and (nw_val or 0) > 0 else '-'
            w_display = str(w_val) if len(space_working_calendars) > 0 and (w_val or 0) > 0 else '-'
            c0 = table.cell(i + 1, 0)
            c0.text = label
            _style_table_cell(c0, bold=True)
            c1 = table.cell(i + 1, 1)
            c1.text = nw_display
            _style_table_cell(c1)
            c2 = table.cell(i + 1, 2)
            c2.text = w_display
            _style_table_cell(c2)

    # ---------- Detailed data ----------
    def _add_detailed_data_section(self, doc):
        """Add detailed time-series data tables with per-page trend line charts matching Excel layout.

        Two layouts:
        - Without base period: single timeline, subtotal row, per-page trend chart
        - With base period: base period & reporting period side-by-side, dual subtotal rows,
          solid/dashed comparison trend chart
        """
        _ = self._

        reporting_data = self.report['reporting_period']
        timestamps = reporting_data.get('timestamps', [])

        if not timestamps or len(timestamps[0]) == 0:
            return

        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        values = reporting_data.get('values', [])
        subtotals = reporting_data.get('subtotals', [])
        ca_len = len(names)

        rows_per_page = 50

        self._add_heading_styled(doc, _('Detailed Data'), level=1)

        if not self.is_base_period_exists:
            # ----- No base period branch -----
            times = timestamps[0]
            if len(times) == 0:
                return

            num_pages = (len(times) + rows_per_page - 1) // rows_per_page

            for page in range(num_pages):
                start_row = page * rows_per_page
                end_row = min(start_row + rows_per_page, len(times))
                page_rows = end_row - start_row

                col_headers = [_('Datetime')]
                for i in range(ca_len):
                    unit_i = units[i] if (units and i < len(units)) else ''
                    col_headers.append(names[i] + ((' (' + unit_i + ')') if unit_i else ''))

                num_cols = len(col_headers)
                # +1 for subtotal row
                table = doc.add_table(rows=page_rows + 2, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for j, h in enumerate(col_headers):
                    c = table.cell(0, j)
                    c.text = h
                    _style_table_cell(c, is_header=True, bold=True, font_size=8)

                for t_idx in range(page_rows):
                    global_idx = start_row + t_idx
                    r_idx = t_idx + 1
                    c0 = table.cell(r_idx, 0)
                    c0.text = str(times[global_idx])
                    _style_table_cell(c0, font_size=8)
                    for j in range(ca_len):
                        col = j + 1
                        val = round2(values[j][global_idx], 2) \
                            if j < len(values) and global_idx < len(values[j]) else ''
                        c = table.cell(r_idx, col)
                        c.text = str(val) if val != '' else ''
                        _style_table_cell(c, font_size=8)

                subtotal_row_idx = page_rows + 1
                c_sub_lbl = table.cell(subtotal_row_idx, 0)
                c_sub_lbl.text = _('Subtotal')
                _style_table_cell(c_sub_lbl, bold=True, font_size=8)
                for i in range(ca_len):
                    col = i + 1
                    c = table.cell(subtotal_row_idx, col)
                    val = subtotals[i] if (subtotals and i < len(subtotals)) else None
                    c.text = str(round2(val, 2)) if val is not None else ''
                    _style_table_cell(c, bold=True, font_size=8)

                # Per-page combined chart
                fig, ax = plt.subplots(figsize=(6.5, 3.5))
                marker_step = max(1, page_rows // 15)
                plotted_any = False
                for i in range(ca_len):
                    data = values[i] if i < len(values) else []
                    page_data = data[start_row:end_row]
                    if not page_data or len(page_data) == 0:
                        continue
                    color = self.chart_colors[i % len(self.chart_colors)]
                    ax.plot(range(len(page_data)), page_data, linewidth=1.2,
                            color=color, label=names[i],
                            marker='o', markersize=3, markevery=marker_step)
                    plotted_any = True
                step = max(1, page_rows // 8)
                ax.set_xticks(range(0, page_rows, step))
                ax.set_xticklabels(
                    [times[start_row + t][:10] for t in range(0, page_rows, step)],
                    rotation=45, ha='right', fontsize=7)
                ax.set_title(_('Reporting Period Consumption') + ' (' +
                             str(start_row + 1) + '-' + str(end_row) + ')',
                             fontsize=10, fontweight='bold')
                if plotted_any:
                    ax.legend(fontsize=6, loc='upper right', ncol=min(ca_len, 3))
                ax.grid(True, alpha=0.3)
                chart_buf = self._fig_to_bytesio(fig, self.dpi)

                doc.add_paragraph('')
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(6.5))

                if page < num_pages - 1:
                    doc.add_page_break()
        else:
            # ----- With base period branch -----
            base_period_data = self.report['base_period']
            base_timestamps = base_period_data.get('timestamps', [])
            base_values = base_period_data.get('values', [])
            base_subtotals = base_period_data.get('subtotals', [])
            base_names = base_period_data.get('names', [])
            base_units = base_period_data.get('units', [])
            base_ca_len = len(base_names)
            reporting_ca_len = ca_len

            base_times = base_timestamps[0] if base_timestamps else []
            reporting_times = timestamps[0]

            max_len = max(len(base_times), len(reporting_times))
            num_pages = (max_len + rows_per_page - 1) // rows_per_page
            marker_step = max(1, rows_per_page // 15)

            for page in range(num_pages):
                start_row = page * rows_per_page
                end_row = min(start_row + rows_per_page, max_len)
                page_rows = end_row - start_row

                col_headers = [_('Base Period') + ' - ' + _('Datetime')]
                for i in range(base_ca_len):
                    bui = base_units[i] if (base_units and i < len(base_units)) else ''
                    col_headers.append(_('Base Period') + ' - ' + base_names[i] +
                                       ((' (' + bui + ')') if bui else ''))
                col_headers.append(_('Reporting Period') + ' - ' + _('Datetime'))
                for i in range(reporting_ca_len):
                    rui = units[i] if (units and i < len(units)) else ''
                    col_headers.append(_('Reporting Period') + ' - ' + names[i] +
                                       ((' (' + rui + ')') if rui else ''))

                num_cols = len(col_headers)
                table = doc.add_table(rows=page_rows + 2, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for j, h in enumerate(col_headers):
                    c = table.cell(0, j)
                    c.text = h
                    _style_table_cell(c, is_header=True, bold=True, font_size=7)

                for t_idx in range(page_rows):
                    global_idx = start_row + t_idx
                    r_idx = t_idx + 1
                    col = 0
                    c_bt = table.cell(r_idx, col)
                    c_bt.text = base_times[global_idx] if global_idx < len(base_times) else ''
                    _style_table_cell(c_bt, font_size=7)
                    col += 1
                    for j in range(base_ca_len):
                        c = table.cell(r_idx, col)
                        if global_idx < len(base_values[j]) if j < len(base_values) else False:
                            c.text = str(round2(base_values[j][global_idx], 2))
                        else:
                            c.text = ''
                        _style_table_cell(c, font_size=7)
                        col += 1
                    c_rt = table.cell(r_idx, col)
                    c_rt.text = reporting_times[global_idx] if global_idx < len(reporting_times) else ''
                    _style_table_cell(c_rt, font_size=7)
                    col += 1
                    for j in range(reporting_ca_len):
                        c = table.cell(r_idx, col)
                        if global_idx < len(values[j]) if j < len(values) else False:
                            c.text = str(round2(values[j][global_idx], 2))
                        else:
                            c.text = ''
                        _style_table_cell(c, font_size=7)
                        col += 1

                # Subtotal row
                sub_idx = page_rows + 1
                col = 0
                c_s0 = table.cell(sub_idx, col)
                c_s0.text = _('Subtotal')
                _style_table_cell(c_s0, bold=True, font_size=7)
                col += 1
                for i in range(base_ca_len):
                    c = table.cell(sub_idx, col)
                    val = base_subtotals[i] if (base_subtotals and i < len(base_subtotals)) else None
                    c.text = str(round2(val, 2)) if val is not None else ''
                    _style_table_cell(c, bold=True, font_size=7)
                    col += 1
                c_s1 = table.cell(sub_idx, col)
                c_s1.text = _('Subtotal')
                _style_table_cell(c_s1, bold=True, font_size=7)
                col += 1
                for i in range(reporting_ca_len):
                    c = table.cell(sub_idx, col)
                    val = subtotals[i] if (subtotals and i < len(subtotals)) else None
                    c.text = str(round2(val, 2)) if val is not None else ''
                    _style_table_cell(c, bold=True, font_size=7)
                    col += 1

                # Per-page comparison chart
                fig, ax = plt.subplots(figsize=(6.5, 3.5))
                plotted_any = False
                for i in range(reporting_ca_len):
                    r_data = values[i] if i < len(values) else []
                    r_page = r_data[start_row:end_row]
                    if not r_page or len(r_page) == 0:
                        continue
                    color = self.chart_colors[i % len(self.chart_colors)]
                    ax.plot(range(len(r_page)), r_page, linewidth=1.2,
                            color=color, marker='o', markersize=3, markevery=marker_step,
                            label=_('Reporting Period') + ' - ' + names[i])
                    plotted_any = True
                    if i < len(base_values):
                        b_data = base_values[i]
                        b_page = b_data[start_row:end_row]
                        if b_page and len(b_page) > 0:
                            ax.plot(range(len(b_page)), b_page, linewidth=1.2,
                                    color=color, linestyle='--', marker='s', markersize=3,
                                    markevery=marker_step,
                                    label=_('Base Period') + ' - ' +
                                          (base_names[i] if i < len(base_names) else ''))
                            plotted_any = True
                step = max(1, page_rows // 8)
                ax.set_xticks(range(0, page_rows, step))
                xlabels = []
                for t in range(0, page_rows, step):
                    gi = start_row + t
                    xlabels.append(reporting_times[gi][:10] if gi < len(reporting_times) else '')
                ax.set_xticklabels(xlabels, rotation=45, ha='right', fontsize=7)
                ax.set_title(str(start_row + 1) + '-' + str(end_row),
                             fontsize=9, fontweight='bold')
                if plotted_any:
                    ax.legend(fontsize=5, loc='upper right', ncol=2)
                ax.grid(True, alpha=0.3)
                chart_buf = self._fig_to_bytesio(fig, self.dpi)

                doc.add_paragraph('')
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(6.5))

                if page < num_pages - 1:
                    doc.add_page_break()

    # ---------- Parameters ----------
    def _add_parameters_section(self, doc):
        """Add parameter data (sensors/tariffs) tables and filled line charts.

        Filters valid parameters: each parameter must have non-empty timestamps
        and values arrays before rendering a compact 25-row table + trend chart.
        """
        _ = self._

        params = self.report.get('parameters', {})
        if not params or not params.get('names') or not params.get('timestamps'):
            return

        param_names = params.get('names', [])
        timestamps = params.get('timestamps', [])
        values = params.get('values', [])
        units = params.get('units', [])

        all_zero = True
        for ts_list in timestamps:
            if ts_list and len(ts_list) > 0:
                all_zero = False
                break
        if all_zero:
            return

        all_params = list(range(len(param_names)))
        valid_params = []
        for i in all_params:
            if i < len(timestamps) and len(timestamps[i]) > 0:
                if i < len(values) and len(values[i]) > 0:
                    valid_params.append(i)
        if not valid_params:
            return

        doc.add_page_break()
        self._add_heading_styled(doc,self.name + ' ' + _('Parameters'), level=1)

        rows_per_param = 10

        for pi in valid_params:
            name = param_names[pi]
            unit_i = units[pi] if (units and pi < len(units)) else ''
            times = timestamps[pi]
            data = values[pi]
            data_len = len(times)

            display_name = name + ((' (' + unit_i + ')') if unit_i else '')

            tbl_rows = min(rows_per_param, data_len)

            fig, ax = plt.subplots(figsize=(5.0, 2.4))
            marker_step_p = max(1, data_len // 20)
            color = '#5B9BD5'
            ax.plot(range(data_len), data, linewidth=1.2,
                    color=color, marker='o', markersize=3,
                    markevery=marker_step_p, label=name)
            ax.fill_between(range(data_len), data, alpha=0.15, color=color)
            step = max(1, data_len // 8)
            ax.set_xticks(range(0, data_len, step))
            ax.set_xticklabels([times[t][:10] for t in range(0, data_len, step)],
                               rotation=45, ha='right', fontsize=6)
            ax.set_ylabel(name, fontsize=8)
            ax.set_title(display_name, fontsize=9, fontweight='bold')
            ax.grid(True, alpha=0.3)
            chart_buf = self._fig_to_bytesio(fig, self.dpi)

            container = doc.add_table(rows=1, cols=2)
            container.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(container)
            left_cell = container.cell(0, 0)
            right_cell = container.cell(0, 1)

            data_table = left_cell.add_table(rows=tbl_rows + 1, cols=2)
            data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            h0 = data_table.cell(0, 0)
            h0.text = _('Time')
            _style_table_cell(h0, is_header=True, bold=True, font_size=8)
            h1 = data_table.cell(0, 1)
            h1.text = name
            _style_table_cell(h1, is_header=True, bold=True, font_size=8)
            for j in range(tbl_rows):
                c_t = data_table.cell(j + 1, 0)
                c_t.text = str(times[j])
                _style_table_cell(c_t, font_size=7)
                c_v = data_table.cell(j + 1, 1)
                c_v.text = str(round2(data[j], 2))
                _style_table_cell(c_v, font_size=7)

            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_buf, width=Inches(4.5))

    # ---------- Base period existence ----------
    def _is_base_period_timestamp_exists(self, base_period_data: Dict) -> bool:
        timestamps = base_period_data.get('timestamps', [])

        if not timestamps:
            return False

        for timestamp in timestamps:
            if timestamp and len(timestamp) > 0:
                return True

        return False


def export(report,
           name,
           base_period_start_datetime_local,
           base_period_end_datetime_local,
           reporting_start_datetime_local,
           reporting_end_datetime_local,
           period_type,
           language):
    """
    Export report data to DOCX and return base64 encoded string.
    This function maintains the same interface as the Excel exporter.   
    """
    exporter = SpaceEnergyDOCXExporter(language)
    return exporter.export(report, name,
                           base_period_start_datetime_local,
                           base_period_end_datetime_local,
                           reporting_start_datetime_local,
                           reporting_end_datetime_local,
                           period_type,
                           language)
