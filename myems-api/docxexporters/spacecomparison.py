"""
Space Comparison DOCX Exporter

This module provides functionality to export space comparison data to DOCX format.
It generates comprehensive reports showing energy consumption comparison between
two spaces with detailed breakdown including difference analysis.

Key Features:
- Space energy comparison analysis (two spaces, single energy category)
- Consumption summary for both spaces
- Detailed data table with difference column
- Comparison line chart (both spaces overlaid)
- Parameter data pages (if available)
- Multi-language support
- Base64 encoding for file transmission

The exported DOCX file includes:
- Cover page with logo and report metadata
- Combined analysis section (consumption summary + comparison chart)
- Detailed data section (paginated table + line chart)
- Parameter data section
"""

import base64
import os
import time
import uuid
import io

from decimal import Decimal
from typing import Optional, Dict, List, Any, BinaryIO
import logging

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.utilities import get_translation, round2

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

_font_setup_done = False


def setup_chinese_fonts():
    global _font_setup_done
    if _font_setup_done:
        return True

    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             '..', 'pdfexporters', 'fonts', 'NotoSansCJK-Regular.ttc')
    try:
        import matplotlib.font_manager as fm
        fm.fontManager.addfont(font_path)
        prop = fm.FontProperties(fname=font_path)
        font_name = prop.get_name()
        plt.rcParams['font.sans-serif'] = [font_name]
        plt.rcParams['axes.unicode_minus'] = False
        logger.info(f"Successfully loaded bundled font: {font_name} from {font_path}")
        _font_setup_done = True
        return True
    except Exception as e:
        logger.warning(f"Failed to load bundled font from {font_path}: {e}")

    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    logger.warning("Failed to load bundled NotoSansCJK font, using DejaVu Sans")
    return False


def _convert_decimals(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    elif isinstance(obj, dict):
        return {k: _convert_decimals(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_convert_decimals(item) for item in obj]
    elif isinstance(obj, tuple):
        return tuple(_convert_decimals(item) for item in obj)
    return obj


def _set_min_row_height(cell, height_twips=150, exact=False):
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def _reset_cell_paragraph_spacing(cell, font_size=9):
    line_twips = max(int(font_size * 20 * 1.15), 120)
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(spacing)
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


def _set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'start', 'bottom', 'end']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _style_table_cell(cell, is_header=False, is_green=False, bold=False, font_size=9):
    cell.paragraphs[0].runs[0].font.bold = bold
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _reset_cell_paragraph_spacing(cell, font_size=font_size)
    _set_cell_margins_zero(cell)
    row_h = max(int(font_size * 20 * 1.3), 160)
    _set_min_row_height(cell, height_twips=row_h, exact=False)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    if is_header:
        shd.set(qn('w:fill'), 'D9E2F3')
    elif is_green:
        shd.set(qn('w:fill'), '90EE90')
    else:
        return
    tcPr.append(shd)


class SpaceComparisonDOCXExporter:
    """
    Export space comparison data to DOCX format.
    Generates comprehensive reports with charts and tables matching Excel layout.
    """

    def __init__(self, language: str = 'zh_CN'):
        font_setup_success = setup_chinese_fonts()
        if not font_setup_success:
            logger.warning("Chinese font setup failed, some text may not display correctly")

        self.language = language
        self.trans = get_translation(language)
        self._ = self.trans.gettext

        self.dpi = 120

        self.chart_colors = ['#4472C4', '#ED7D31', '#70AD47', '#FFC000', '#5B9BD5',
                             '#FF6B6B', '#9B59B6', '#1ABC9C', '#E67E22', '#2ECC71',
                             '#3498DB', '#E74C3C', '#2ECC71', '#F39C12', '#9B59B6']

    def export(self,
               report: Dict[str, Any],
               space1_name: str,
               space2_name: str,
               energy_category_name: str,
               reporting_start_datetime_local: str,
               reporting_end_datetime_local: str,
               period_type: str,
               language: str) -> Optional[str]:
        if report is None:
            return None
        start_time = time.time()
        logger.info(f"Starting DOCX generation for SpaceComparison: {space1_name} vs {space2_name}")

        docx_filename = self.generate_docx(
            report, space1_name, space2_name, energy_category_name,
            reporting_start_datetime_local, reporting_end_datetime_local,
            period_type, language
        )

        result = ''
        if docx_filename and os.path.exists(docx_filename):
            try:
                with open(docx_filename, 'rb') as binary_file:
                    binary_data = binary_file.read()
                result = base64.b64encode(binary_data).decode('utf-8')
            except Exception as e:
                logger.error(f"Failed to encode DOCX: {str(e)}")
            finally:
                try:
                    os.remove(docx_filename)
                except Exception:
                    pass
        elapsed = time.time() - start_time
        logger.info(f"DOCX generation completed in {elapsed:.2f}s for SpaceComparison")
        return result

    @staticmethod
    def _fig_to_bytesio(fig, dpi: int) -> BinaryIO:
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    @staticmethod
    def _sanitize_values(values):
        result = []
        for v in values:
            if v is None:
                result.append(float('nan'))
            else:
                try:
                    result.append(float(v))
                except (TypeError, ValueError):
                    result.append(float('nan'))
        return result

    @staticmethod
    def _filter_valid_data(data):
        xs, ys = [], []
        for idx, v in enumerate(data):
            if v is not None:
                try:
                    ys.append(float(v))
                    xs.append(idx)
                except (TypeError, ValueError):
                    pass
        return xs, ys

    def generate_docx(self,
                      report: Dict[str, Any],
                      space1_name: str,
                      space2_name: str,
                      energy_category_name: str,
                      reporting_start_datetime_local: str,
                      reporting_end_datetime_local: str,
                      period_type: str,
                      language: str) -> Optional[str]:
        _ = self._

        if "reporting_period1" not in report.keys() or \
                "values" not in report['reporting_period1'].keys() or \
                len(report['reporting_period1']['values']) == 0:
            doc = Document()
            section = doc.sections[0]
            section.orientation = 1
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            self._add_cover_page(doc, space1_name, space2_name,
                                 energy_category_name,
                                 reporting_start_datetime_local,
                                 reporting_end_datetime_local,
                                 period_type)
            filename = str(uuid.uuid4()) + '.docx'
            doc.save(filename)
            return filename

        filename = str(uuid.uuid4()) + '.docx'

        self.report = _convert_decimals(report)
        self.space1_name = space1_name
        self.space2_name = space2_name
        self.energy_category_name = energy_category_name
        self.reporting_start = reporting_start_datetime_local
        self.reporting_end = reporting_end_datetime_local
        self.period_type = period_type
        self.unit = report['energy_category']['unit_of_measure']

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        self._add_cover_page(doc, space1_name, space2_name,
                             energy_category_name,
                             reporting_start_datetime_local,
                             reporting_end_datetime_local,
                             period_type)

        self._add_combined_analysis_section(doc)
        self._add_detailed_data_section(doc)
        self._add_parameters_section(doc)

        doc.save(filename)
        logger.info(f"DOCX generated: {filename}")
        return filename

    def _add_heading_styled(self, doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        return heading

    def _add_cover_page(self, doc, space1_name, space2_name,
                        energy_category_name,
                        reporting_start, reporting_end,
                        period_type):
        _ = self._
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                '..', 'excelexporters', 'myems.png')
        if os.path.exists(img_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(7.0))
            except Exception as e:
                logger.warning(f"Failed to load logo image: {e}")

        for _unused in range(3):
            doc.add_paragraph('')

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(_('Space Comparison Analysis'))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for _unused in range(3):
            doc.add_paragraph('')

        info_data = [
            [_('Space') + '1:', space1_name],
            [_('Space') + '2:', space2_name],
            [_('Energy Category') + ':', energy_category_name],
            [_('Period Type') + ':', period_type],
            [_('Reporting Start Datetime') + ':', reporting_start],
            [_('Reporting End Datetime') + ':', reporting_end],
        ]

        info_table = doc.add_table(rows=len(info_data), cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(info_table)

        half_w = Inches(3.2)
        for i, (label, value) in enumerate(info_data):
            c_label = info_table.cell(i, 0)
            c_label.width = half_w
            c_label.text = label
            c_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = c_label.paragraphs[0].runs[0]
            r_lbl.bold = True
            r_lbl.font.size = Pt(13)
            r_lbl.font.name = 'Arial'
            r_lbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            c_value = info_table.cell(i, 1)
            c_value.width = half_w
            c_value.text = str(value) if value is not None else ''
            c_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_val = c_value.paragraphs[0].runs[0]
            r_val.font.size = Pt(13)
            r_val.font.name = 'Arial'
            r_val._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        doc.add_page_break()

    def _add_combined_analysis_section(self, doc):
        """Add combined analysis section: Consumption summary table + comparison chart."""
        _ = self._
        reporting_data1 = self.report['reporting_period1']
        reporting_data2 = self.report['reporting_period2']
        diff_data = self.report['diff']

        total1 = round2(reporting_data1.get('total_in_category', 0), 2)
        total2 = round2(reporting_data2.get('total_in_category', 0), 2)
        total_diff = round2(diff_data.get('total_in_category', 0), 2)

        unit = self.unit
        cat_name = self.energy_category_name

        self._add_heading_styled(doc, self.space1_name + ' & ' + self.space2_name + ' - ' +
                                 _('Reporting Period Consumption'), level=1)

        summary_data = [
            ['', cat_name + ' (' + unit + ')'],
            [self.space1_name, str(total1)],
            [self.space2_name, str(total2)],
            [_('Difference'), str(total_diff)],
        ]

        table = doc.add_table(rows=len(summary_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(summary_data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:
                    _style_table_cell(cell, is_green=True, bold=True)
                elif j == 0:
                    _style_table_cell(cell, is_green=True, bold=True)
                else:
                    _style_table_cell(cell)

        doc.add_paragraph('')

        timestamps1 = reporting_data1.get('timestamps', [])
        values1 = reporting_data1.get('values', [])
        values2 = reporting_data2.get('values', [])

        if timestamps1 and len(timestamps1) > 0:
            xs = list(range(len(timestamps1)))
            ys1 = self._sanitize_values(values1)
            ys2 = self._sanitize_values(values2)

            fig, ax = plt.subplots(figsize=(9.75, 5.25))
            marker_step = max(1, len(xs) // 30)

            ax.plot(xs, ys1, linewidth=1.5, color='#4472C4',
                    marker='o', markersize=3,
                    markevery=marker_step,
                    label=self.space1_name)
            ax.plot(xs, ys2, linewidth=1.5, color='#ED7D31',
                    marker='s', markersize=3,
                    markevery=marker_step,
                    label=self.space2_name)

            step = max(1, len(timestamps1) // 10)
            ax.set_xticks(range(0, len(timestamps1), step))
            ax.set_xticklabels(
                [timestamps1[t][:10] for t in range(0, len(timestamps1), step)],
                rotation=45, ha='right', fontsize=7)

            ax.set_ylabel(cat_name + ' (' + unit + ')', fontsize=9)
            ax.legend(fontsize=9)
            ax.grid(True, alpha=0.3)

            chart_buf = self._fig_to_bytesio(fig, self.dpi)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_buf, width=Inches(9.75))

        doc.add_page_break()

    def _add_detailed_data_section(self, doc):
        """Add detailed time-series data tables ."""
        _ = self._

        reporting_data1 = self.report['reporting_period1']
        reporting_data2 = self.report['reporting_period2']
        diff_data = self.report['diff']

        timestamps = reporting_data1.get('timestamps', [])
        values1 = reporting_data1.get('values', [])
        values2 = reporting_data2.get('values', [])
        diff_values = diff_data.get('values', [])

        if not timestamps or len(timestamps) == 0:
            return

        unit = self.unit
        cat_name = self.energy_category_name
        rows_per_page = 50

        self._add_heading_styled(doc, self.space1_name + ' and ' + self.space2_name + ' ' +
                                 _('Detailed Data'), level=1)

        num_pages = (len(timestamps) + rows_per_page - 1) // rows_per_page

        for page in range(num_pages):
            start_row = page * rows_per_page
            end_row = min(start_row + rows_per_page, len(timestamps))
            page_rows = end_row - start_row

            col_headers = [
                _('Datetime'),
                self.space1_name + ' ' + cat_name + ' (' + unit + ')',
                self.space2_name + ' ' + cat_name + ' (' + unit + ')',
                _('Difference')
            ]

            num_cols = len(col_headers)
            table = doc.add_table(rows=page_rows + 2, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(col_headers):
                c = table.cell(0, j)
                c.text = h
                _style_table_cell(c, is_header=True, bold=True, font_size=8)

            for t_idx in range(page_rows):
                global_idx = start_row + t_idx
                r_idx = t_idx + 1
                c0 = table.cell(r_idx, 0)
                c0.text = str(timestamps[global_idx])
                _style_table_cell(c0, font_size=8)

                v1 = round2(values1[global_idx], 2) \
                    if global_idx < len(values1) and values1[global_idx] is not None else ''
                c1 = table.cell(r_idx, 1)
                c1.text = str(v1) if v1 != '' else ''
                _style_table_cell(c1, font_size=8)

                v2 = round2(values2[global_idx], 2) \
                    if global_idx < len(values2) and values2[global_idx] is not None else ''
                c2 = table.cell(r_idx, 2)
                c2.text = str(v2) if v2 != '' else ''
                _style_table_cell(c2, font_size=8)

                vd = round2(diff_values[global_idx], 2) \
                    if global_idx < len(diff_values) and diff_values[global_idx] is not None else ''
                c3 = table.cell(r_idx, 3)
                c3.text = str(vd) if vd != '' else ''
                _style_table_cell(c3, font_size=8)

            total_row_idx = page_rows + 1
            total1 = round2(reporting_data1.get('total_in_category', 0), 2)
            total2 = round2(reporting_data2.get('total_in_category', 0), 2)
            total_diff = round2(diff_data.get('total_in_category', 0), 2)

            c_t0 = table.cell(total_row_idx, 0)
            c_t0.text = _('Total')
            _style_table_cell(c_t0, bold=True, font_size=8)
            c_t1 = table.cell(total_row_idx, 1)
            c_t1.text = str(total1)
            _style_table_cell(c_t1, bold=True, font_size=8)
            c_t2 = table.cell(total_row_idx, 2)
            c_t2.text = str(total2)
            _style_table_cell(c_t2, bold=True, font_size=8)
            c_t3 = table.cell(total_row_idx, 3)
            c_t3.text = str(total_diff)
            _style_table_cell(c_t3, bold=True, font_size=8)

            if page < num_pages - 1:
                doc.add_page_break()

    def _add_parameters_section(self, doc):
        """Add parameter data sections for both space1 and space2."""
        _ = self._

        for param_key, space_name in [('parameters1', self.space1_name),
                                       ('parameters2', self.space2_name)]:
            params = self.report.get(param_key, {})
            if not params or not params.get('names') or not params.get('timestamps'):
                continue

            param_names = params.get('names', [])
            timestamps = params.get('timestamps', [])
            values = params.get('values', [])
            units = params.get('units', [])

            all_zero = True
            for ts_list in timestamps:
                if ts_list and len(ts_list) > 0:
                    all_zero = False
                    break
            if all_zero:
                continue

            valid_params = []
            for i in range(len(param_names)):
                if i < len(timestamps) and len(timestamps[i]) > 0:
                    if i < len(values) and len(values[i]) > 0:
                        valid_params.append(i)
            if not valid_params:
                continue

            self._add_heading_styled(doc, space_name + ' ' + _('Parameters'), level=1)

            rows_per_param = 10

            for pi in valid_params:
                name = param_names[pi]
                unit_i = units[pi] if (units and pi < len(units)) else ''
                times = timestamps[pi]
                data = values[pi]
                data_len = len(times)

                display_name = name + ((' (' + unit_i + ')') if unit_i else '')

                tbl_rows = min(rows_per_param, data_len)

                fig, ax = plt.subplots(figsize=(5.0, 2.4))
                xs, ys = self._filter_valid_data(data)
                if ys:
                    marker_step_p = max(1, data_len // 20)
                    color = '#5B9BD5'
                    ax.plot(range(data_len), data, linewidth=1.2,
                            color=color, marker='o', markersize=3,
                            markevery=marker_step_p, label=name)
                    ax.fill_between(range(data_len), data, alpha=0.15, color=color)
                step = max(1, data_len // 8)
                ax.set_xticks(range(0, data_len, step))
                ax.set_xticklabels([times[t][:10] for t in range(0, data_len, step)],
                                   rotation=45, ha='right', fontsize=6)
                ax.set_ylabel(name, fontsize=8)
                ax.set_title(display_name, fontsize=9, fontweight='bold')
                ax.grid(True, alpha=0.3)
                chart_buf = self._fig_to_bytesio(fig, self.dpi)

                container = doc.add_table(rows=1, cols=2)
                container.alignment = WD_TABLE_ALIGNMENT.CENTER
                _remove_table_borders(container)
                left_cell = container.cell(0, 0)
                right_cell = container.cell(0, 1)

                data_table = left_cell.add_table(rows=tbl_rows + 1, cols=2)
                data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                h0 = data_table.cell(0, 0)
                h0.text = _('Time')
                _style_table_cell(h0, is_header=True, bold=True, font_size=8)
                h1 = data_table.cell(0, 1)
                h1.text = name
                _style_table_cell(h1, is_header=True, bold=True, font_size=8)
                for j in range(tbl_rows):
                    c_t = data_table.cell(j + 1, 0)
                    c_t.text = str(times[j])
                    _style_table_cell(c_t, font_size=7)
                    c_v = data_table.cell(j + 1, 1)
                    val = round2(data[j], 2) if data[j] is not None else ''
                    c_v.text = str(val) if val != '' else ''
                    _style_table_cell(c_v, font_size=7)

                p = right_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(chart_buf, width=Inches(4.5))


def export(report,
           name,
           base_period_start_datetime_local,
           base_period_end_datetime_local,
           reporting_start_datetime_local,
           reporting_end_datetime_local,
           period_type,
           language):
    """
    Export report data to DOCX and return base64 encoded string.
    Provides a unified signature matching other DOCX exporters.
    Extracts space1_name, space2_name, and energy_category_name from report data.
    """
    space1_name = report.get('space1', {}).get('name', name)
    space2_name = report.get('space2', {}).get('name', '')
    energy_category_name = report.get('energy_category', {}).get('name', '')
    exporter = SpaceComparisonDOCXExporter(language)
    return exporter.export(report, space1_name, space2_name, energy_category_name,
                           reporting_start_datetime_local,
                           reporting_end_datetime_local,
                           period_type, language)
