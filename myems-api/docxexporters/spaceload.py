"""
Space Load DOCX Exporter

This module provides functionality to export space load data to DOCX format.
It generates comprehensive reports showing load analysis for spaces
with detailed breakdown by energy categories and time periods.

Key Features:
- Space load analysis (average load, maximum load, load factor)
- Base period vs reporting period comparison
- Load breakdown by energy categories
- Detailed data with line charts
- Multi-language support
- Base64 encoding for file transmission

The exported DOCX file includes:
- Cover page with logo and report metadata
- Combined analysis page (3 tables: average load, maximum load, load factor)
- Detailed data charts (each category has average + maximum load charts)
- Parameter data pages
"""

import base64
import os
import time
import uuid
import io

from decimal import Decimal
from typing import Optional, Dict, List, Any, BinaryIO
import logging

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.utilities import get_translation, round2

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

_font_setup_done = False


def setup_chinese_fonts():
    global _font_setup_done
    if _font_setup_done:
        return True

    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             '..', 'pdfexporters', 'fonts', 'NotoSansCJK-Regular.ttc')
    try:
        import matplotlib.font_manager as fm
        fm.fontManager.addfont(font_path)
        prop = fm.FontProperties(fname=font_path)
        font_name = prop.get_name()
        plt.rcParams['font.sans-serif'] = [font_name]
        plt.rcParams['axes.unicode_minus'] = False
        logger.info(f"Successfully loaded bundled font: {font_name} from {font_path}")
        _font_setup_done = True
        return True
    except Exception as e:
        logger.warning(f"Failed to load bundled font from {font_path}: {e}")

    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    logger.warning("Failed to load bundled NotoSansCJK font, using DejaVu Sans")
    return False


def _convert_decimals(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    elif isinstance(obj, dict):
        return {k: _convert_decimals(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_convert_decimals(item) for item in obj]
    elif isinstance(obj, tuple):
        return tuple(_convert_decimals(item) for item in obj)
    return obj


def _set_min_row_height(cell, height_twips=150, exact=False):
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def _reset_cell_paragraph_spacing(cell, font_size=9):
    line_twips = max(int(font_size * 20 * 1.15), 120)
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(spacing)
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


def _set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'start', 'bottom', 'end']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _style_table_cell(cell, is_header=False, is_green=False, bold=False, font_size=9):
    cell.paragraphs[0].runs[0].font.bold = bold
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _reset_cell_paragraph_spacing(cell, font_size=font_size)
    _set_cell_margins_zero(cell)
    row_h = max(int(font_size * 20 * 1.3), 160)
    _set_min_row_height(cell, height_twips=row_h, exact=False)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    if is_header:
        shd.set(qn('w:fill'), 'D9E2F3')
    elif is_green:
        shd.set(qn('w:fill'), '90EE90')
    else:
        return
    tcPr.append(shd)


class SpaceLoadDOCXExporter:
    """
    Export space load data to DOCX format.
    Generates comprehensive reports with charts and tables matching Excel layout.
    """

    def __init__(self, language: str = 'zh_CN'):
        font_setup_success = setup_chinese_fonts()
        if not font_setup_success:
            logger.warning("Chinese font setup failed, some text may not display correctly")

        self.language = language
        self.trans = get_translation(language)
        self._ = self.trans.gettext

        self.dpi = 120

        self.chart_colors = ['#4472C4', '#ED7D31', '#70AD47', '#FFC000', '#5B9BD5',
                             '#FF6B6B', '#9B59B6', '#1ABC9C', '#E67E22', '#2ECC71',
                             '#3498DB', '#E74C3C', '#2ECC71', '#F39C12', '#9B59B6']

    def export(self,
               report: Dict[str, Any],
               name: str,
               base_period_start_datetime_local: str,
               base_period_end_datetime_local: str,
               reporting_start_datetime_local: str,
               reporting_end_datetime_local: str,
               period_type: str,
               language: str) -> Optional[str]:
        if report is None:
            return None
        start_time = time.time()
        logger.info(f"Starting DOCX generation for {name}")

        docx_filename = self.generate_docx(
            report, name,
            base_period_start_datetime_local,
            base_period_end_datetime_local,
            reporting_start_datetime_local,
            reporting_end_datetime_local,
            period_type,
            language
        )

        result = ''
        if docx_filename and os.path.exists(docx_filename):
            try:
                with open(docx_filename, 'rb') as binary_file:
                    binary_data = binary_file.read()
                result = base64.b64encode(binary_data).decode('utf-8')
            except Exception as e:
                logger.error(f"Failed to encode DOCX: {str(e)}")
            finally:
                try:
                    os.remove(docx_filename)
                except Exception:
                    pass
        elapsed = time.time() - start_time
        logger.info(f"DOCX generation completed in {elapsed:.2f}s for {name}")
        return result

    @staticmethod
    def _fig_to_bytesio(fig, dpi: int) -> BinaryIO:
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    @staticmethod
    def _filter_valid_data(data):
        xs, ys = [], []
        for idx, v in enumerate(data):
            if v is not None:
                try:
                    ys.append(float(v))
                    xs.append(idx)
                except (TypeError, ValueError):
                    pass
        return xs, ys

    def generate_docx(self,
                      report: Dict[str, Any],
                      name: str,
                      base_period_start_datetime_local: str,
                      base_period_end_datetime_local: str,
                      reporting_start_datetime_local: str,
                      reporting_end_datetime_local: str,
                      period_type: str,
                      language: str) -> Optional[str]:
        _ = self._

        if "reporting_period" not in report.keys() or \
                "names" not in report['reporting_period'].keys() or \
                len(report['reporting_period']['names']) == 0:
            doc = Document()
            section = doc.sections[0]
            section.orientation = 1
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            self._add_cover_page(doc, name, period_type,
                                 reporting_start_datetime_local,
                                 reporting_end_datetime_local,
                                 base_period_start_datetime_local,
                                 base_period_end_datetime_local,
                                 False)
            filename = str(uuid.uuid4()) + '.docx'
            doc.save(filename)
            return filename

        filename = str(uuid.uuid4()) + '.docx'

        self.report = _convert_decimals(report)
        self.name = name
        self.base_period_start = base_period_start_datetime_local
        self.base_period_end = base_period_end_datetime_local
        self.reporting_start = reporting_start_datetime_local
        self.reporting_end = reporting_end_datetime_local
        self.period_type = period_type

        self.is_base_period_exists = self._is_base_period_timestamp_exists(report['base_period'])

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        self._add_cover_page(doc, name, period_type,
                             reporting_start_datetime_local,
                             reporting_end_datetime_local,
                             base_period_start_datetime_local,
                             base_period_end_datetime_local,
                             self.is_base_period_exists)

        self._add_combined_analysis(doc)
        self._add_detailed_data_charts(doc)
        self._add_parameters_section(doc)

        doc.save(filename)
        logger.info(f"DOCX generated: {filename}")
        return filename

    def _add_heading_styled(self, doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        return heading

    def _add_cover_page(self, doc, name, period_type,
                        reporting_start, reporting_end,
                        base_period_start, base_period_end,
                        has_base_period):
        _ = self._
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                '..', 'excelexporters', 'myems.png')
        if os.path.exists(img_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(7.0))
            except Exception as e:
                logger.warning(f"Failed to load logo image: {e}")

        for _unused in range(3):
            doc.add_paragraph('')

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(_('Load'))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for _unused in range(3):
            doc.add_paragraph('')

        info_data = [
            [_('Name') + ':', name],
            [_('Period Type') + ':', period_type],
            [_('Reporting Start Datetime') + ':', reporting_start],
            [_('Reporting End Datetime') + ':', reporting_end],
        ]
        if has_base_period:
            info_data.append([_('Base Period Start Datetime') + ':', base_period_start])
            info_data.append([_('Base Period End Datetime') + ':', base_period_end])

        info_table = doc.add_table(rows=len(info_data), cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(info_table)

        half_w = Inches(3.2)
        for i, (label, value) in enumerate(info_data):
            c_label = info_table.cell(i, 0)
            c_label.width = half_w
            c_label.text = label
            c_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = c_label.paragraphs[0].runs[0]
            r_lbl.bold = True
            r_lbl.font.size = Pt(13)
            r_lbl.font.name = 'Arial'
            r_lbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            c_value = info_table.cell(i, 1)
            c_value.width = half_w
            c_value.text = str(value) if value is not None else ''
            c_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_val = c_value.paragraphs[0].runs[0]
            r_val.font.size = Pt(13)
            r_val.font.name = 'Arial'
            r_val._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        doc.add_page_break()

    def _add_combined_analysis(self, doc):
        _ = self._
        reporting_data = self.report['reporting_period']
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        averages = reporting_data.get('averages', [])
        averages_per_unit_area = reporting_data.get('averages_per_unit_area', [])
        averages_increment_rate = reporting_data.get('averages_increment_rate', [])
        maximums = reporting_data.get('maximums', [])
        maximums_per_unit_area = reporting_data.get('maximums_per_unit_area', [])
        maximums_increment_rate = reporting_data.get('maximums_increment_rate', [])
        factors = reporting_data.get('factors', [])
        factors_increment_rate = reporting_data.get('factors_increment_rate', [])
        ca_len = len(names)

        if ca_len == 0:
            return

        col_headers_with_units = ['']
        for i in range(ca_len):
            unit_i = units[i] if (units and i < len(units)) else ''
            col_headers_with_units.append(names[i] + ((' (' + unit_i + '/H)') if unit_i else ''))

        col_headers_names_only = ['']
        for i in range(ca_len):
            col_headers_names_only.append(names[i])

        self._add_heading_styled(doc, self.name + _('Reporting Period Average Load'), level=1)

        num_cols = ca_len + 1
        table1 = doc.add_table(rows=4, cols=num_cols)
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(col_headers_with_units):
            cell = table1.cell(0, j)
            cell.text = h
            _style_table_cell(cell, is_header=True, bold=True)

        row_labels_avg = [_('Average Load'), _('Per Unit Area'), _('Increment Rate')]
        for r_idx, row_label in enumerate(row_labels_avg, start=1):
            cell = table1.cell(r_idx, 0)
            cell.text = row_label
            _style_table_cell(cell, is_green=True, bold=True)

        for i in range(ca_len):
            col = i + 1
            c1 = table1.cell(1, col)
            val = averages[i] if averages and i < len(averages) else None
            c1.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(c1)

            c2 = table1.cell(2, col)
            val = averages_per_unit_area[i] if averages_per_unit_area and i < len(averages_per_unit_area) else None
            c2.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(c2)

            c3 = table1.cell(3, col)
            val = averages_increment_rate[i] if averages_increment_rate and i < len(averages_increment_rate) else None
            c3.text = (str(round2(val * 100, 2)) + '%') if val is not None else ''
            _style_table_cell(c3)

        doc.add_paragraph('')

        self._add_heading_styled(doc, self.name + _('Reporting Period Maximum Load'), level=1)

        table2 = doc.add_table(rows=4, cols=num_cols)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(col_headers_with_units):
            cell = table2.cell(0, j)
            cell.text = h
            _style_table_cell(cell, is_header=True, bold=True)

        row_labels_max = [_('Maximum Load'), _('Per Unit Area'), _('Increment Rate')]
        for r_idx, row_label in enumerate(row_labels_max, start=1):
            cell = table2.cell(r_idx, 0)
            cell.text = row_label
            _style_table_cell(cell, is_green=True, bold=True)

        for i in range(ca_len):
            col = i + 1
            c1 = table2.cell(1, col)
            val = maximums[i] if maximums and i < len(maximums) else None
            c1.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(c1)

            c2 = table2.cell(2, col)
            val = maximums_per_unit_area[i] if maximums_per_unit_area and i < len(maximums_per_unit_area) else None
            c2.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(c2)

            c3 = table2.cell(3, col)
            val = maximums_increment_rate[i] if maximums_increment_rate and i < len(maximums_increment_rate) else None
            c3.text = (str(round2(val * 100, 2)) + '%') if val is not None else ''
            _style_table_cell(c3)

        doc.add_paragraph('')

        self._add_heading_styled(doc, self.name + _('Reporting Period Load Factor'), level=1)

        table3 = doc.add_table(rows=3, cols=num_cols)
        table3.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(col_headers_names_only):
            cell = table3.cell(0, j)
            cell.text = h
            _style_table_cell(cell, is_header=True, bold=True)

        row_labels_factor = [_('Load Factor'), _('Increment Rate')]
        for r_idx, row_label in enumerate(row_labels_factor, start=1):
            cell = table3.cell(r_idx, 0)
            cell.text = row_label
            _style_table_cell(cell, is_green=True, bold=True)

        for i in range(ca_len):
            col = i + 1
            c1 = table3.cell(1, col)
            val = factors[i] if factors and i < len(factors) else None
            c1.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(c1)

            c2 = table3.cell(2, col)
            val = factors_increment_rate[i] if factors_increment_rate and i < len(factors_increment_rate) else None
            c2.text = (str(round2(val * 100, 2)) + '%') if val is not None else ''
            _style_table_cell(c2)

        doc.add_paragraph('')

    def _add_detailed_data_charts(self, doc):
        _ = self._

        reporting_data = self.report['reporting_period']
        timestamps = reporting_data.get('timestamps', [])
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        sub_averages = reporting_data.get('sub_averages', [])
        sub_maximums = reporting_data.get('sub_maximums', [])

        if not timestamps or len(timestamps[0]) == 0 or not names:
            return

        reporting_times = timestamps[0]
        num_categories = len(names)

        has_sub_averages = len(sub_averages) > 0
        has_sub_maximums = len(sub_maximums) > 0

        charts_per_category = 0
        if has_sub_averages:
            charts_per_category += 1
        if has_sub_maximums:
            charts_per_category += 1

        if charts_per_category == 0:
            return

        doc.add_page_break()
        self._add_heading_styled(doc, self.name + ' ' + _('Detailed Data'), level=1)

        fig_w, fig_h = 4.8, 3.0
        display_w = 4.8

        all_charts = []

        def _set_ticks(ax, raw_len):
            step = max(1, raw_len // 10)
            ax.set_xticks(range(0, raw_len, step))
            ax.set_xticklabels(
                [reporting_times[t][:10] if t < len(reporting_times) else ''
                 for t in range(0, raw_len, step)],
                rotation=45, ha='right', fontsize=7)

        if not self.is_base_period_exists:
            for i in range(num_categories):
                color = self.chart_colors[i % len(self.chart_colors)]
                unit_i = units[i] if (units and i < len(units)) else ''

                if has_sub_averages:
                    raw_data = sub_averages[i] if i < len(sub_averages) else []
                    xs, ys = self._filter_valid_data(raw_data)

                    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
                    marker_step = max(1, len(ys) // 30) if ys else 1
                    if ys:
                        ax.plot(xs, ys, linewidth=1.2, color=color,
                                marker='o', markersize=3,
                                markevery=marker_step)
                    _set_ticks(ax, len(raw_data))
                    ax.set_title(_('Reporting Period Average Load') + ' - ' +
                                 names[i] + ' ' + _('Average Load') +
                                 ((' (' + unit_i + '/H)') if unit_i else ''),
                                 fontsize=9, fontweight='bold')
                    ax.grid(True, alpha=0.3)
                    all_charts.append(self._fig_to_bytesio(fig, self.dpi))

                if has_sub_maximums:
                    raw_data = sub_maximums[i] if i < len(sub_maximums) else []
                    xs, ys = self._filter_valid_data(raw_data)

                    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
                    marker_step = max(1, len(ys) // 30) if ys else 1
                    if ys:
                        ax.plot(xs, ys, linewidth=1.2, color=color,
                                marker='o', markersize=3,
                                markevery=marker_step)
                    _set_ticks(ax, len(raw_data))
                    ax.set_title(_('Reporting Period Maximum Load') + ' - ' +
                                 names[i] + ' ' + _('Maximum Load') +
                                 ((' (' + unit_i + '/H)') if unit_i else ''),
                                 fontsize=9, fontweight='bold')
                    ax.grid(True, alpha=0.3)
                    all_charts.append(self._fig_to_bytesio(fig, self.dpi))
        else:
            base_period_data = self.report['base_period']
            base_sub_averages = base_period_data.get('sub_averages', [])
            base_sub_maximums = base_period_data.get('sub_maximums', [])

            for i in range(num_categories):
                color = self.chart_colors[i % len(self.chart_colors)]
                unit_i = units[i] if (units and i < len(units)) else ''

                if has_sub_averages:
                    fig, ax = plt.subplots(figsize=(fig_w, fig_h))

                    r_data = sub_averages[i] if i < len(sub_averages) else []
                    r_xs, r_ys = self._filter_valid_data(r_data)
                    marker_step = max(1, len(r_ys) // 30) if r_ys else 1
                    if r_ys:
                        ax.plot(r_xs, r_ys, linewidth=1.2, color=color,
                                marker='o', markersize=3,
                                markevery=marker_step,
                                label=_('Reporting Period') + ' - ' + names[i])

                    if i < len(base_sub_averages):
                        b_data = base_sub_averages[i]
                        b_xs, b_ys = self._filter_valid_data(b_data)
                        marker_step_b = max(1, len(b_ys) // 30) if b_ys else 1
                        if b_ys:
                            ax.plot(b_xs, b_ys, linewidth=1.2, color=color,
                                    linestyle='--', marker='s', markersize=3,
                                    markevery=marker_step_b,
                                    label=_('Base Period') + ' - ' + names[i])

                    _set_ticks(ax, len(r_data))
                    ax.set_title(
                        _('Base Period Average Load') + ' / ' +
                        _('Reporting Period Average Load') + ' - ' +
                        names[i] + ' ' + _('Average Load') +
                        ((' (' + unit_i + '/H)') if unit_i else ''),
                        fontsize=8, fontweight='bold')
                    if len(r_ys) > 0 or (i < len(base_sub_averages) and len(
                            self._filter_valid_data(base_sub_averages[i])[1]) > 0):
                        ax.legend(fontsize=7)
                    ax.grid(True, alpha=0.3)
                    all_charts.append(self._fig_to_bytesio(fig, self.dpi))

                if has_sub_maximums:
                    fig, ax = plt.subplots(figsize=(fig_w, fig_h))

                    r_data = sub_maximums[i] if i < len(sub_maximums) else []
                    r_xs, r_ys = self._filter_valid_data(r_data)
                    marker_step = max(1, len(r_ys) // 30) if r_ys else 1
                    if r_ys:
                        ax.plot(r_xs, r_ys, linewidth=1.2, color=color,
                                marker='o', markersize=3,
                                markevery=marker_step,
                                label=_('Reporting Period') + ' - ' + names[i])

                    if i < len(base_sub_maximums):
                        b_data = base_sub_maximums[i]
                        b_xs, b_ys = self._filter_valid_data(b_data)
                        marker_step_b = max(1, len(b_ys) // 30) if b_ys else 1
                        if b_ys:
                            ax.plot(b_xs, b_ys, linewidth=1.2, color=color,
                                    linestyle='--', marker='s', markersize=3,
                                    markevery=marker_step_b,
                                    label=_('Base Period') + ' - ' + names[i])

                    _set_ticks(ax, len(r_data))
                    ax.set_title(
                        _('Base Period Maximum Load') + ' / ' +
                        _('Reporting Period Maximum Load') + ' - ' +
                        names[i] + ' ' + _('Maximum Load') +
                        ((' (' + unit_i + '/H)') if unit_i else ''),
                        fontsize=8, fontweight='bold')
                    if len(r_ys) > 0 or (i < len(base_sub_maximums) and len(
                            self._filter_valid_data(base_sub_maximums[i])[1]) > 0):
                        ax.legend(fontsize=7)
                    ax.grid(True, alpha=0.3)
                    all_charts.append(self._fig_to_bytesio(fig, self.dpi))

        num_total_charts = len(all_charts)
        charts_per_page = 4

        for page_start in range(0, num_total_charts, charts_per_page):
            page_end = min(page_start + charts_per_page, num_total_charts)
            page_indices = list(range(page_start, page_end))
            num_on_page = len(page_indices)

            rows = (num_on_page + 1) // 2

            for row_idx in range(rows):
                slot0 = row_idx * 2
                slot1 = slot0 + 1
                has_left = slot0 < num_on_page
                has_right = slot1 < num_on_page

                if has_left and not has_right:
                    container = doc.add_table(rows=1, cols=2)
                    container.alignment = WD_TABLE_ALIGNMENT.CENTER
                    _remove_table_borders(container)
                    container.cell(0, 0).merge(container.cell(0, 1))
                    cell = container.cell(0, 0)

                    chart_buf = all_charts[page_indices[slot0]]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(chart_buf, width=Inches(display_w))
                else:
                    container_cols = min(2, num_on_page - row_idx * 2)
                    container = doc.add_table(rows=1, cols=container_cols)
                    container.alignment = WD_TABLE_ALIGNMENT.CENTER
                    _remove_table_borders(container)

                    for col_idx in range(container_cols):
                        slot = row_idx * 2 + col_idx
                        if slot >= num_on_page:
                            continue
                        chart_buf = all_charts[page_indices[slot]]

                        cell = container.cell(0, col_idx)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(chart_buf, width=Inches(display_w))

            if page_end < num_total_charts:
                doc.add_page_break()

        doc.add_page_break()

    def _add_parameters_section(self, doc):
        _ = self._

        params = self.report.get('parameters', {})
        if not params or not params.get('names') or not params.get('timestamps'):
            return

        param_names = params.get('names', [])
        timestamps = params.get('timestamps', [])
        values = params.get('values', [])
        units = params.get('units', [])

        all_zero = True
        for ts_list in timestamps:
            if ts_list and len(ts_list) > 0:
                all_zero = False
                break
        if all_zero:
            return

        all_params = list(range(len(param_names)))
        valid_params = []
        for i in all_params:
            if i < len(timestamps) and len(timestamps[i]) > 0:
                if i < len(values) and len(values[i]) > 0:
                    valid_params.append(i)
        if not valid_params:
            return

        self._add_heading_styled(doc, self.name + ' ' + _('Parameters'), level=1)

        rows_per_param = 10

        for pi in valid_params:
            name = param_names[pi]
            unit_i = units[pi] if (units and pi < len(units)) else ''
            times = timestamps[pi]
            data = values[pi]
            data_len = len(times)

            display_name = name + ((' (' + unit_i + ')') if unit_i else '')

            tbl_rows = min(rows_per_param, data_len)

            fig, ax = plt.subplots(figsize=(5.0, 2.4))
            marker_step_p = max(1, data_len // 20)
            color = '#5B9BD5'
            ax.plot(range(data_len), data, linewidth=1.2,
                    color=color, marker='o', markersize=3,
                    markevery=marker_step_p, label=name)
            ax.fill_between(range(data_len), data, alpha=0.15, color=color)
            step = max(1, data_len // 8)
            ax.set_xticks(range(0, data_len, step))
            ax.set_xticklabels([times[t][:10] for t in range(0, data_len, step)],
                               rotation=45, ha='right', fontsize=6)
            ax.set_ylabel(name, fontsize=8)
            ax.set_title(display_name, fontsize=9, fontweight='bold')
            ax.grid(True, alpha=0.3)
            chart_buf = self._fig_to_bytesio(fig, self.dpi)

            container = doc.add_table(rows=1, cols=2)
            container.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(container)
            left_cell = container.cell(0, 0)
            right_cell = container.cell(0, 1)

            data_table = left_cell.add_table(rows=tbl_rows + 1, cols=2)
            data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            h0 = data_table.cell(0, 0)
            h0.text = _('Time')
            _style_table_cell(h0, is_header=True, bold=True, font_size=8)
            h1 = data_table.cell(0, 1)
            h1.text = name
            _style_table_cell(h1, is_header=True, bold=True, font_size=8)
            for j in range(tbl_rows):
                c_t = data_table.cell(j + 1, 0)
                c_t.text = str(times[j])
                _style_table_cell(c_t, font_size=7)
                c_v = data_table.cell(j + 1, 1)
                c_v.text = str(round2(data[j], 2))
                _style_table_cell(c_v, font_size=7)

            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_buf, width=Inches(4.5))

    def _is_base_period_timestamp_exists(self, base_period_data: Dict) -> bool:
        timestamps = base_period_data.get('timestamps', [])

        if not timestamps:
            return False

        for timestamp in timestamps:
            if timestamp and len(timestamp) > 0:
                return True

        return False


def export(report,
           name,
           base_period_start_datetime_local,
           base_period_end_datetime_local,
           reporting_start_datetime_local,
           reporting_end_datetime_local,
           period_type,
           language):
    """
    Export report data to DOCX and return base64 encoded string.
    This function maintains the same interface as the Excel exporter.
    """
    exporter = SpaceLoadDOCXExporter(language)
    return exporter.export(report, name,
                           base_period_start_datetime_local,
                           base_period_end_datetime_local,
                           reporting_start_datetime_local,
                           reporting_end_datetime_local,
                           period_type,
                           language)
