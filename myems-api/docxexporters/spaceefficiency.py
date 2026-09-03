"""
Space Efficiency DOCX Exporter

This module provides functionality to export space efficiency data to DOCX format.
It generates comprehensive reports showing efficiency analysis for spaces
with detailed breakdown by energy categories and time periods.

Key Features:
- Space efficiency analysis
- Base period vs reporting period comparison
- Efficiency breakdown by energy categories
- Detailed data with line charts
- Multi-language support
- Base64 encoding for file transmission

The exported DOCX file includes:
- Cover page with logo and report metadata
- Combined analysis section (reporting period cumulative efficiency table)
- Detailed data charts (paginated, up to 4 per page in 2x2 grid)
- Parameter data section
"""

import base64
import os
import time
import uuid
import io

from decimal import Decimal
from typing import Optional, Dict, List, Any, BinaryIO
import logging

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.utilities import get_translation, round2

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

_font_setup_done = False


def setup_chinese_fonts():
    global _font_setup_done
    if _font_setup_done:
        return True

    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             '..', 'pdfexporters', 'fonts', 'NotoSansCJK-Regular.ttc')
    try:
        import matplotlib.font_manager as fm
        fm.fontManager.addfont(font_path)
        prop = fm.FontProperties(fname=font_path)
        font_name = prop.get_name()
        plt.rcParams['font.sans-serif'] = [font_name]
        plt.rcParams['axes.unicode_minus'] = False
        logger.info(f"Successfully loaded bundled font: {font_name} from {font_path}")
        _font_setup_done = True
        return True
    except Exception as e:
        logger.warning(f"Failed to load bundled font from {font_path}: {e}")

    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    logger.warning("Failed to load bundled NotoSansCJK font, using DejaVu Sans")
    return False


def _convert_decimals(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    elif isinstance(obj, dict):
        return {k: _convert_decimals(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_convert_decimals(item) for item in obj]
    elif isinstance(obj, tuple):
        return tuple(_convert_decimals(item) for item in obj)
    return obj


def _set_min_row_height(cell, height_twips=150, exact=False):
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def _reset_cell_paragraph_spacing(cell, font_size=9):
    line_twips = max(int(font_size * 20 * 1.15), 120)
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(spacing)
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


def _set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'start', 'bottom', 'end']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _style_table_cell(cell, is_header=False, is_green=False, bold=False, font_size=9):
    cell.paragraphs[0].runs[0].font.bold = bold
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _reset_cell_paragraph_spacing(cell, font_size=font_size)
    _set_cell_margins_zero(cell)
    row_h = max(int(font_size * 20 * 1.3), 160)
    _set_min_row_height(cell, height_twips=row_h, exact=False)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    if is_header:
        shd.set(qn('w:fill'), 'D9E2F3')
    elif is_green:
        shd.set(qn('w:fill'), '90EE90')
    else:
        return
    tcPr.append(shd)


class SpaceEfficiencyDOCXExporter:
    """
    Export space efficiency data to DOCX format.
    Generates comprehensive reports with charts and tables matching Excel layout.
    """

    def __init__(self, language: str = 'zh_CN'):
        font_setup_success = setup_chinese_fonts()
        if not font_setup_success:
            logger.warning("Chinese font setup failed, some text may not display correctly")

        self.language = language
        self.trans = get_translation(language)
        self._ = self.trans.gettext

        self.dpi = 120

        self.chart_colors = ['#4472C4', '#ED7D31', '#70AD47', '#FFC000', '#5B9BD5',
                             '#FF6B6B', '#9B59B6', '#1ABC9C', '#E67E22', '#2ECC71',
                             '#3498DB', '#E74C3C', '#2ECC71', '#F39C12', '#9B59B6']

    def export(self,
               report: Dict[str, Any],
               name: str,
               base_period_start_datetime_local: str,
               base_period_end_datetime_local: str,
               reporting_start_datetime_local: str,
               reporting_end_datetime_local: str,
               period_type: str,
               language: str) -> Optional[str]:
        if report is None:
            return None
        start_time = time.time()
        logger.info(f"Starting DOCX generation for {name}")

        docx_filename = self.generate_docx(
            report, name,
            base_period_start_datetime_local,
            base_period_end_datetime_local,
            reporting_start_datetime_local,
            reporting_end_datetime_local,
            period_type,
            language
        )

        result = ''
        if docx_filename and os.path.exists(docx_filename):
            try:
                with open(docx_filename, 'rb') as binary_file:
                    binary_data = binary_file.read()
                result = base64.b64encode(binary_data).decode('utf-8')
            except Exception as e:
                logger.error(f"Failed to encode DOCX: {str(e)}")
            finally:
                try:
                    os.remove(docx_filename)
                except Exception:
                    pass
        elapsed = time.time() - start_time
        logger.info(f"DOCX generation completed in {elapsed:.2f}s for {name}")
        return result

    @staticmethod
    def _fig_to_bytesio(fig, dpi: int) -> BinaryIO:
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    @staticmethod
    def _filter_valid_data(data):
        xs, ys = [], []
        for idx, v in enumerate(data):
            if v is not None:
                try:
                    ys.append(float(v))
                    xs.append(idx)
                except (TypeError, ValueError):
                    pass
        return xs, ys

    def generate_docx(self,
                      report: Dict[str, Any],
                      name: str,
                      base_period_start_datetime_local: str,
                      base_period_end_datetime_local: str,
                      reporting_start_datetime_local: str,
                      reporting_end_datetime_local: str,
                      period_type: str,
                      language: str) -> Optional[str]:
        _ = self._

        if "reporting_period_efficiency" not in report.keys() or \
                "names" not in report['reporting_period_efficiency'].keys() or \
                len(report['reporting_period_efficiency']['names']) == 0:
            doc = Document()
            section = doc.sections[0]
            section.orientation = 1
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            self._add_cover_section(doc, name, period_type,
                                    reporting_start_datetime_local,
                                    reporting_end_datetime_local,
                                    base_period_start_datetime_local,
                                    base_period_end_datetime_local,
                                    False)
            filename = str(uuid.uuid4()) + '.docx'
            doc.save(filename)
            return filename

        filename = str(uuid.uuid4()) + '.docx'

        self.report = _convert_decimals(report)
        self.name = name
        self.base_period_start = base_period_start_datetime_local
        self.base_period_end = base_period_end_datetime_local
        self.reporting_start = reporting_start_datetime_local
        self.reporting_end = reporting_end_datetime_local
        self.period_type = period_type

        base_data = report.get('base_period_efficiency', {})
        self.is_base_period_exists = self._is_base_period_timestamp_exists(base_data)

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        self._add_cover_section(doc, name, period_type,
                                reporting_start_datetime_local,
                                reporting_end_datetime_local,
                                base_period_start_datetime_local,
                                base_period_end_datetime_local,
                                self.is_base_period_exists)

        self._add_combined_analysis_section(doc)
        self._add_detailed_data_charts_section(doc)
        self._add_parameters_section(doc)

        doc.save(filename)
        logger.info(f"DOCX generated: {filename}")
        return filename

    def _add_heading_styled(self, doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        return heading

    def _add_cover_section(self, doc, name, period_type,
                           reporting_start, reporting_end,
                           base_period_start, base_period_end,
                           has_base_period):
        _ = self._
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                '..', 'excelexporters', 'myems.png')
        if os.path.exists(img_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(7.0))
            except Exception as e:
                logger.warning(f"Failed to load logo image: {e}")

        for _unused in range(3):
            doc.add_paragraph('')

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(_('Efficiency'))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for _unused in range(3):
            doc.add_paragraph('')

        info_data = [
            [_('Name') + ':', name],
            [_('Period Type') + ':', period_type],
            [_('Reporting Start Datetime') + ':', reporting_start],
            [_('Reporting End Datetime') + ':', reporting_end],
        ]
        if has_base_period:
            info_data.append([_('Base Period Start Datetime') + ':', base_period_start])
            info_data.append([_('Base Period End Datetime') + ':', base_period_end])

        info_table = doc.add_table(rows=len(info_data), cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(info_table)

        half_w = Inches(3.2)
        for i, (label, value) in enumerate(info_data):
            c_label = info_table.cell(i, 0)
            c_label.width = half_w
            c_label.text = label
            c_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = c_label.paragraphs[0].runs[0]
            r_lbl.bold = True
            r_lbl.font.size = Pt(13)
            r_lbl.font.name = 'Arial'
            r_lbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            c_value = info_table.cell(i, 1)
            c_value.width = half_w
            c_value.text = str(value) if value is not None else ''
            c_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_val = c_value.paragraphs[0].runs[0]
            r_val.font.size = Pt(13)
            r_val.font.name = 'Arial'
            r_val._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        doc.add_page_break()

    def _add_combined_analysis_section(self, doc):
        _ = self._
        reporting_data = self.report['reporting_period_efficiency']
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        cumulations = reporting_data.get('cumulations', [])
        increment_rates = reporting_data.get('increment_rates', [])
        ca_len = len(names)

        if ca_len == 0:
            return

        self._add_heading_styled(doc, self.name + ' - ' + _('Reporting Period Cumulative Efficiency'), level=1)

        num_cols = ca_len + 1
        table = doc.add_table(rows=3, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['']
        for i in range(ca_len):
            unit_i = units[i] if (units and i < len(units)) else ''
            headers.append(names[i] + ((' (' + unit_i + ')') if unit_i else ''))

        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
            _style_table_cell(cell, is_header=True, bold=True)

        row_labels = [_('Cumulative Efficiency'), _('Increment Rate')]
        for r_idx, row_label in enumerate(row_labels, start=1):
            cell = table.cell(r_idx, 0)
            cell.text = row_label
            _style_table_cell(cell, is_green=True, bold=True)

        for i in range(ca_len):
            col = i + 1
            cell_cum = table.cell(1, col)
            val = cumulations[i] if cumulations and i < len(cumulations) else None
            cell_cum.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(cell_cum)

            cell_inc = table.cell(2, col)
            val = increment_rates[i] if increment_rates and i < len(increment_rates) else None
            cell_inc.text = (str(round2(val * 100, 2)) + '%') if val is not None else '-'
            _style_table_cell(cell_inc)

        doc.add_paragraph('')

    def _add_detailed_data_charts_section(self, doc):
        _ = self._

        reporting_data = self.report['reporting_period_efficiency']
        timestamps = reporting_data.get('timestamps', [])
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        values = reporting_data.get('values', [])

        if not timestamps or len(timestamps[0]) == 0 or not names:
            return

        reporting_times = timestamps[0]
        num_categories = len(names)

        doc.add_page_break()
        self._add_heading_styled(doc, self.name + ' ' + _('Detailed Data'), level=1)

        chart_colors = self.chart_colors
        N = len(chart_colors)

        if not self.is_base_period_exists:
            per_row_charts = 2
            max_per_page = 4
            drawn_so_far = 0

            for row_start in range(0, num_categories, per_row_charts):
                row_end = min(row_start + per_row_charts, num_categories)
                row_categories = list(range(row_start, row_end))
                row_size = len(row_categories)

                container = doc.add_table(rows=1, cols=per_row_charts)
                container.alignment = WD_TABLE_ALIGNMENT.CENTER
                _remove_table_borders(container)

                if row_size == 1:
                    container.cell(0, 0).merge(container.cell(0, 1))

                for local_idx, i in enumerate(row_categories):
                    raw_data = values[i] if i < len(values) else []
                    unit_i = units[i] if (units and i < len(units)) else ''
                    color = chart_colors[i % N]
                    data_len = len(raw_data)
                    marker_step = max(1, data_len // 30)
                    safe_data = [(v if isinstance(v, (int, float)) else 0)
                                 if v is not None else 0 for v in raw_data]

                    fig_w, fig_h = 4.8, 3.0
                    single_mode = (row_size == 1)
                    if single_mode:
                        fig_w, fig_h = 4.8, 3.0

                    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
                    if len(safe_data) > 0:
                        ax.plot(range(len(safe_data)), safe_data,
                                linewidth=1.2,
                                color=color, marker='o', markersize=3,
                                markevery=marker_step)
                    step = max(1, data_len // 10) if data_len > 0 else 1
                    ax.set_xticks(range(0, data_len, step))
                    xlabels = []
                    for t in range(0, data_len, step):
                        xlabels.append(
                            reporting_times[t][:10] if t < len(reporting_times) else '')
                    ax.set_xticklabels(xlabels, rotation=45, ha='right', fontsize=7)
                    title = _('Reporting Period Cumulative Efficiency') + ' - ' + \
                            names[i] + ((' (' + unit_i + ')') if unit_i else '')
                    ax.set_title(title, fontsize=9, fontweight='bold')
                    ax.grid(True, alpha=0.3)
                    plt.tight_layout(pad=1.0)
                    chart_buf = self._fig_to_bytesio(fig, self.dpi)

                    cell_idx = 0 if single_mode else local_idx
                    cell = container.cell(0, cell_idx)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(chart_buf, width=Inches(4.8))

                drawn_so_far += row_size
                if drawn_so_far >= max_per_page and drawn_so_far < num_categories \
                        and (drawn_so_far % max_per_page == 0):
                    doc.add_page_break()
        else:
            base_period_data = self.report['base_period_efficiency']
            base_timestamps = base_period_data.get('timestamps', [])
            base_values = base_period_data.get('values', [])
            base_names = base_period_data.get('names', [])
            base_units = base_period_data.get('units', [])

            per_row_charts = 2
            max_per_page = 4
            drawn_so_far = 0

            for row_start in range(0, num_categories, per_row_charts):
                row_end = min(row_start + per_row_charts, num_categories)
                row_categories = list(range(row_start, row_end))
                row_size = len(row_categories)

                container = doc.add_table(rows=1, cols=per_row_charts)
                container.alignment = WD_TABLE_ALIGNMENT.CENTER
                _remove_table_borders(container)

                if row_size == 1:
                    container.cell(0, 0).merge(container.cell(0, 1))

                for local_idx, i in enumerate(row_categories):
                    r_raw = values[i] if i < len(values) else []
                    unit_i = units[i] if (units and i < len(units)) else ''
                    b_raw = base_values[i] if i < len(base_values) else None
                    b_name = base_names[i] if i < len(base_names) else names[i]
                    color = chart_colors[i % N]
                    data_len = len(r_raw)
                    marker_step = max(1, data_len // 30)

                    safe_r = [(v if isinstance(v, (int, float)) else 0)
                              if v is not None else 0 for v in r_raw]
                    if b_raw is not None:
                        safe_b = [(v if isinstance(v, (int, float)) else 0)
                                  if v is not None else 0 for v in b_raw]
                    else:
                        safe_b = None

                    fig_w, fig_h = 4.8, 3.0
                    single_mode = (row_size == 1)
                    if single_mode:
                        fig_w, fig_h = 4.8, 3.0

                    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
                    has_reporting_line = (len(safe_r) > 0)
                    has_base_line = (safe_b is not None and len(safe_b) > 0)
                    if has_reporting_line:
                        ax.plot(range(len(safe_r)), safe_r,
                                linewidth=1.2,
                                color=color, marker='o', markersize=3,
                                markevery=marker_step,
                                label=_('Reporting Period') + ' - ' + names[i])
                    if has_base_line:
                        safe_b_cut = safe_b[:len(safe_r)]
                        x_b = list(range(len(safe_b_cut)))
                        ax.plot(x_b, safe_b_cut,
                                linewidth=1.2,
                                color=color, linestyle='--', marker='s', markersize=3,
                                markevery=marker_step,
                                label=_('Base Period') + ' - ' + b_name)

                    step = max(1, data_len // 10) if data_len > 0 else 1
                    ax.set_xticks(range(0, data_len, step))
                    xlabels = []
                    for t in range(0, data_len, step):
                        xlabels.append(
                            reporting_times[t][:10] if t < len(reporting_times) else '')
                    ax.set_xticklabels(xlabels, rotation=45, ha='right', fontsize=7)
                    title = _('Base Period Efficiency') + ' / ' + \
                            _('Reporting Period Efficiency') + ' - ' + \
                            names[i] + ((' (' + unit_i + ')') if unit_i else '')
                    ax.set_title(title, fontsize=9, fontweight='bold')
                    if has_reporting_line or has_base_line:
                        ax.legend(fontsize=7)
                    ax.grid(True, alpha=0.3)
                    plt.tight_layout(pad=1.0)
                    chart_buf = self._fig_to_bytesio(fig, self.dpi)

                    cell_idx = 0 if single_mode else local_idx
                    cell = container.cell(0, cell_idx)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(chart_buf, width=Inches(4.8))

                drawn_so_far += row_size
                if drawn_so_far >= max_per_page and drawn_so_far < num_categories \
                        and (drawn_so_far % max_per_page == 0):
                    doc.add_page_break()

    def _add_parameters_section(self, doc):
        _ = self._

        params = self.report.get('parameters', {})
        if not params or not params.get('names') or not params.get('timestamps'):
            return

        param_names = params.get('names', [])
        timestamps = params.get('timestamps', [])
        values = params.get('values', [])
        units = params.get('units', [])

        all_zero = True
        for ts_list in timestamps:
            if ts_list and len(ts_list) > 0:
                all_zero = False
                break
        if all_zero:
            return

        all_params = list(range(len(param_names)))
        valid_params = []
        for i in all_params:
            if i < len(timestamps) and len(timestamps[i]) > 0:
                if i < len(values) and len(values[i]) > 0:
                    valid_params.append(i)
        if not valid_params:
            return

        doc.add_page_break()
        self._add_heading_styled(doc, self.name + ' ' + _('Parameters'), level=1)

        rows_per_param = 10

        for pi in valid_params:
            name = param_names[pi]
            unit_i = units[pi] if (units and pi < len(units)) else ''
            times = timestamps[pi]
            data = values[pi]
            data_len = len(times)

            display_name = name + ((' (' + unit_i + ')') if unit_i else '')

            tbl_rows = min(rows_per_param, data_len)

            fig, ax = plt.subplots(figsize=(5.0, 2.4))
            marker_step_p = max(1, data_len // 20)
            color = '#5B9BD5'
            ax.plot(range(data_len), data, linewidth=1.2,
                    color=color, marker='o', markersize=3,
                    markevery=marker_step_p, label=name)
            ax.fill_between(range(data_len), data, alpha=0.15, color=color)
            step = max(1, data_len // 8)
            ax.set_xticks(range(0, data_len, step))
            ax.set_xticklabels([times[t][:10] for t in range(0, data_len, step)],
                               rotation=45, ha='right', fontsize=6)
            ax.set_ylabel(name, fontsize=8)
            ax.set_title(display_name, fontsize=9, fontweight='bold')
            ax.grid(True, alpha=0.3)
            chart_buf = self._fig_to_bytesio(fig, self.dpi)

            container = doc.add_table(rows=1, cols=2)
            container.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(container)
            left_cell = container.cell(0, 0)
            right_cell = container.cell(0, 1)

            data_table = left_cell.add_table(rows=tbl_rows + 1, cols=2)
            data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            h0 = data_table.cell(0, 0)
            h0.text = _('Time')
            _style_table_cell(h0, is_header=True, bold=True, font_size=8)
            h1 = data_table.cell(0, 1)
            h1.text = name
            _style_table_cell(h1, is_header=True, bold=True, font_size=8)
            for j in range(tbl_rows):
                c_t = data_table.cell(j + 1, 0)
                c_t.text = str(times[j])
                _style_table_cell(c_t, font_size=7)
                c_v = data_table.cell(j + 1, 1)
                c_v.text = str(round2(data[j], 2))
                _style_table_cell(c_v, font_size=7)

            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_buf, width=Inches(4.5))

    def _is_base_period_timestamp_exists(self, base_period_data: Dict) -> bool:
        timestamps = base_period_data.get('timestamps', [])

        if not timestamps:
            return False

        for timestamp in timestamps:
            if timestamp and len(timestamp) > 0:
                return True

        return False


def export(report,
           name,
           base_period_start_datetime_local,
           base_period_end_datetime_local,
           reporting_start_datetime_local,
           reporting_end_datetime_local,
           period_type,
           language):
    """
    Export report data to DOCX and return base64 encoded string.
    This function maintains the same interface as the Excel exporter.
    """
    exporter = SpaceEfficiencyDOCXExporter(language)
    return exporter.export(report, name,
                        base_period_start_datetime_local,
                        base_period_end_datetime_local,
                        reporting_start_datetime_local,
                        reporting_end_datetime_local,
                        period_type,
                        language)
