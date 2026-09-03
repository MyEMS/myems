"""
Space Carbon DOCX Exporter

This module provides functionality to export space carbon emissions data to DOCX format.
It generates comprehensive reports showing carbon dioxide emissions for spaces
with detailed analysis including base period comparison, time-of-use breakdown,
and proportion analysis.

Key Features:
- Space carbon emissions analysis
- Base period vs reporting period comparison
- Time-of-use electricity carbon breakdown
- Carbon emissions proportion charts
- Detailed data tables with separate line charts per energy category
- Multi-language support
- Base64 encoding for file transmission

The exported DOCX file includes:
- Cover page with logo and report metadata
- Combined analysis section (reporting period table + time-of-use + carbon proportion)
- Detailed data section (paginated tables)
- Separate line charts for each energy category
- Parameter data tables with filled line charts
"""

import base64
import os
import time
import uuid
import io

from decimal import Decimal
from typing import Optional, Dict, List, Any, BinaryIO
import logging

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.utilities import get_translation, round2

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

_font_setup_done = False


def setup_chinese_fonts():
    global _font_setup_done
    if _font_setup_done:
        return True

    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             '..', 'pdfexporters', 'fonts', 'NotoSansCJK-Regular.ttc')
    try:
        import matplotlib.font_manager as fm
        fm.fontManager.addfont(font_path)
        prop = fm.FontProperties(fname=font_path)
        font_name = prop.get_name()
        plt.rcParams['font.sans-serif'] = [font_name]
        plt.rcParams['axes.unicode_minus'] = False
        logger.info(f"Successfully loaded bundled font: {font_name} from {font_path}")
        _font_setup_done = True
        return True
    except Exception as e:
        logger.warning(f"Failed to load bundled font from {font_path}: {e}")

    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    logger.warning("Failed to load bundled NotoSansCJK font, using DejaVu Sans")
    return False


def _convert_decimals(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    elif isinstance(obj, dict):
        return {k: _convert_decimals(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_convert_decimals(item) for item in obj]
    elif isinstance(obj, tuple):
        return tuple(_convert_decimals(item) for item in obj)
    return obj


def _set_min_row_height(cell, height_twips=150, exact=False):
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def _reset_cell_paragraph_spacing(cell, font_size=9):
    line_twips = max(int(font_size * 20 * 1.15), 120)
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(spacing)
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


def _set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'start', 'bottom', 'end']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _style_table_cell(cell, is_header=False, is_green=False, bold=False, font_size=9):
    cell.paragraphs[0].runs[0].font.bold = bold
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _reset_cell_paragraph_spacing(cell, font_size=font_size)
    _set_cell_margins_zero(cell)
    row_h = max(int(font_size * 20 * 1.3), 160)
    _set_min_row_height(cell, height_twips=row_h, exact=False)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    if is_header:
        shd.set(qn('w:fill'), 'D9E2F3')
    elif is_green:
        shd.set(qn('w:fill'), '90EE90')
    else:
        return
    tcPr.append(shd)


class SpaceCarbonDOCXExporter:
    """
    Export space carbon emissions data to DOCX format.
    Generates comprehensive reports with charts and tables matching Excel layout.
    """

    def __init__(self, language: str = 'zh_CN'):
        font_setup_success = setup_chinese_fonts()
        if not font_setup_success:
            logger.warning("Chinese font setup failed, some text may not display correctly")

        self.language = language
        self.trans = get_translation(language)
        self._ = self.trans.gettext

        self.dpi = 120

        self.chart_colors = ['#4472C4', '#ED7D31', '#70AD47', '#FFC000', '#5B9BD5',
                             '#FF6B6B', '#9B59B6', '#1ABC9C', '#E67E22', '#2ECC71',
                             '#3498DB', '#E74C3C', '#2ECC71', '#F39C12', '#9B59B6']

    def export(self,
               report: Dict[str, Any],
               name: str,
               base_period_start_datetime_local: str,
               base_period_end_datetime_local: str,
               reporting_start_datetime_local: str,
               reporting_end_datetime_local: str,
               period_type: str,
               language: str) -> Optional[str]:
        if report is None:
            return None
        start_time = time.time()
        logger.info(f"Starting DOCX generation for {name}")

        docx_filename = self.generate_docx(
            report, name,
            base_period_start_datetime_local,
            base_period_end_datetime_local,
            reporting_start_datetime_local,
            reporting_end_datetime_local,
            period_type,
            language
        )

        result = ''
        if docx_filename and os.path.exists(docx_filename):
            try:
                with open(docx_filename, 'rb') as binary_file:
                    binary_data = binary_file.read()
                result = base64.b64encode(binary_data).decode('utf-8')
            except Exception as e:
                logger.error(f"Failed to encode DOCX: {str(e)}")
            finally:
                try:
                    os.remove(docx_filename)
                except Exception:
                    pass
        elapsed = time.time() - start_time
        logger.info(f"DOCX generation completed in {elapsed:.2f}s for {name}")
        return result

    @staticmethod
    def _fig_to_bytesio(fig, dpi: int) -> BinaryIO:
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    def _make_pie_chart(self, values, labels, title, colors=None):
        if not values or sum((v or 0) for v in values) == 0:
            return None
        fig, ax = plt.subplots(figsize=(3.2, 2.6))
        if colors is None:
            colors = self.chart_colors[:len(labels)]
        filtered = [(l, v, c) for l, v, c in zip(labels, values, colors) if (v or 0) > 0]
        if not filtered:
            plt.close(fig)
            return None
        f_labels, f_values, f_colors = zip(*filtered)
        if len(f_labels) > 8:
            sorted_data = sorted(zip(f_labels, f_values, f_colors), key=lambda x: x[1], reverse=True)
            top_data = sorted_data[:7]
            other_sum = sum(v for _, v, _ in sorted_data[7:])
            f_labels = [l for l, _, _ in top_data] + [self._('Others')]
            f_values = [v for _, v, _ in top_data] + [other_sum]
            f_colors = [c for _, _, c in top_data] + ['#999999']
        ax.pie(f_values, labels=f_labels, autopct='%1.1f%%', colors=f_colors, startangle=90)
        ax.set_title(title, fontsize=10, fontweight='bold')
        return self._fig_to_bytesio(fig, self.dpi)

    def generate_docx(self,
                      report: Dict[str, Any],
                      name: str,
                      base_period_start_datetime_local: str,
                      base_period_end_datetime_local: str,
                      reporting_start_datetime_local: str,
                      reporting_end_datetime_local: str,
                      period_type: str,
                      language: str) -> Optional[str]:
        _ = self._

        if "reporting_period" not in report.keys() or \
                "names" not in report['reporting_period'].keys() or \
                len(report['reporting_period']['names']) == 0:
            doc = Document()
            section = doc.sections[0]
            section.orientation = 1
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            self._add_cover_page(doc, name, period_type,
                                 reporting_start_datetime_local,
                                 reporting_end_datetime_local,
                                 base_period_start_datetime_local,
                                 base_period_end_datetime_local,
                                 False)
            filename = str(uuid.uuid4()) + '.docx'
            doc.save(filename)
            return filename

        filename = str(uuid.uuid4()) + '.docx'

        self.report = _convert_decimals(report)
        self.name = name
        self.base_period_start = base_period_start_datetime_local
        self.base_period_end = base_period_end_datetime_local
        self.reporting_start = reporting_start_datetime_local
        self.reporting_end = reporting_end_datetime_local
        self.period_type = period_type

        self.is_base_period_exists = self._is_base_period_timestamp_exists(report['base_period'])

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        self._add_cover_page(doc, name, period_type,
                             reporting_start_datetime_local,
                             reporting_end_datetime_local,
                             base_period_start_datetime_local,
                             base_period_end_datetime_local,
                             self.is_base_period_exists)

        self._add_combined_analysis_section(doc)
        self._add_detailed_data_section(doc)
        self._add_detailed_data_charts_section(doc)
        self._add_parameters_section(doc)

        doc.save(filename)
        logger.info(f"DOCX generated: {filename}")
        return filename

    # ---------- Utility ----------
    def _add_heading_styled(self, doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        return heading

    def _add_cover_page(self, doc, name, period_type,
                        reporting_start, reporting_end,
                        base_period_start, base_period_end,
                        has_base_period):
        _ = self._
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                '..', 'excelexporters', 'myems.png')
        if os.path.exists(img_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(7.0))
            except Exception as e:
                logger.warning(f"Failed to load logo image: {e}")

        for _unused in range(3):
            doc.add_paragraph('')

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(_('Carbon'))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for _unused in range(3):
            doc.add_paragraph('')

        info_data = [
            [_('Name') + ':', name],
            [_('Period Type') + ':', period_type],
            [_('Reporting Start Datetime') + ':', reporting_start],
            [_('Reporting End Datetime') + ':', reporting_end],
        ]
        if has_base_period:
            info_data.append([_('Base Period Start Datetime') + ':', base_period_start])
            info_data.append([_('Base Period End Datetime') + ':', base_period_end])

        info_table = doc.add_table(rows=len(info_data), cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(info_table)

        half_w = Inches(3.2)
        for i, (label, value) in enumerate(info_data):
            c_label = info_table.cell(i, 0)
            c_label.width = half_w
            c_label.text = label
            c_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = c_label.paragraphs[0].runs[0]
            r_lbl.bold = True
            r_lbl.font.size = Pt(13)
            r_lbl.font.name = 'Arial'
            r_lbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            c_value = info_table.cell(i, 1)
            c_value.width = half_w
            c_value.text = str(value) if value is not None else ''
            c_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_val = c_value.paragraphs[0].runs[0]
            r_val.font.size = Pt(13)
            r_val.font.name = 'Arial'
            r_val._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        doc.add_page_break()

    # ---------- Combined analysis ----------
    def _add_combined_analysis_section(self, doc):
        _ = self._
        reporting_data = self.report['reporting_period']
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        subtotals = reporting_data.get('subtotals', [])
        subtotals_per_unit_area = reporting_data.get('subtotals_per_unit_area', [])
        increment_rates = reporting_data.get('increment_rates', [])
        total_unit = reporting_data.get('total_unit', 'KGCO2E')
        ca_len = len(names)

        if ca_len == 0:
            return

        self._add_heading_styled(doc, self.name + ' - ' + _('Reporting Period Carbon Dioxide Emissions'), level=1)

        num_cols = ca_len + 2
        table = doc.add_table(rows=4, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['']
        for i in range(ca_len):
            unit_i = units[i] if (units and i < len(units)) else ''
            headers.append(names[i] + ((' (' + unit_i + ')') if unit_i else ''))
        headers.append(_('Total') + '(' + total_unit + ')')

        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
            _style_table_cell(cell, is_header=True, bold=True)

        row_labels = [_('Carbon Dioxide Emissions'), _('Per Unit Area'), _('Increment Rate')]
        for r_idx, row_label in enumerate(row_labels, start=1):
            cell = table.cell(r_idx, 0)
            cell.text = row_label
            _style_table_cell(cell, is_green=True, bold=True)

        total_val = reporting_data.get('total', 0)
        total_per_area = reporting_data.get('total_per_unit_area', None)
        total_inc_rate = reporting_data.get('total_increment_rate', None)

        for i in range(ca_len):
            col = i + 1
            cell_cons = table.cell(1, col)
            val = subtotals[i] if (subtotals and i < len(subtotals)) else None
            cell_cons.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(cell_cons)

            cell_area = table.cell(2, col)
            val = subtotals_per_unit_area[i] if (subtotals_per_unit_area and i < len(subtotals_per_unit_area)) else None
            cell_area.text = str(round2(val, 2)) if val is not None else ''
            _style_table_cell(cell_area)

            cell_inc = table.cell(3, col)
            val = increment_rates[i] if (increment_rates and i < len(increment_rates)) else None
            cell_inc.text = (str(round2(val * 100, 2)) + '%') if val is not None else ''
            _style_table_cell(cell_inc)

        total_col = ca_len + 1
        table.cell(1, total_col).text = str(round2(total_val, 2))
        _style_table_cell(table.cell(1, total_col))
        total_area_text = str(round2(total_per_area, 2)) if total_per_area is not None else ''
        table.cell(2, total_col).text = total_area_text
        _style_table_cell(table.cell(2, total_col))
        total_inc_text = (str(round2(total_inc_rate * 100, 2)) + '%') if total_inc_rate is not None else ''
        table.cell(3, total_col).text = total_inc_text
        _style_table_cell(table.cell(3, total_col))

        doc.add_paragraph('')

        electricity_index = -1
        for i in range(len(reporting_data.get('energy_category_ids', []))):
            if reporting_data['energy_category_ids'][i] == 1:
                electricity_index = i
                break

        tou_exists = electricity_index >= 0
        tou_categories = [_('TopPeak'), _('OnPeak'), _('MidPeak'), _('OffPeak')]
        tou_values = []
        if tou_exists:
            toppeaks = reporting_data.get('toppeaks', [])
            onpeaks = reporting_data.get('onpeaks', [])
            midpeaks = reporting_data.get('midpeaks', [])
            offpeaks = reporting_data.get('offpeaks', [])
            tou_values = [
                round2(toppeaks[electricity_index], 2) if electricity_index < len(toppeaks) else 0,
                round2(onpeaks[electricity_index], 2) if electricity_index < len(onpeaks) else 0,
                round2(midpeaks[electricity_index], 2) if electricity_index < len(midpeaks) else 0,
                round2(offpeaks[electricity_index], 2) if electricity_index < len(offpeaks) else 0,
            ]

        carbon_exists = subtotals and sum((v or 0) for v in subtotals) > 0
        carbon_values = [round2(v, 3) for v in subtotals] if carbon_exists else []

        container = doc.add_table(rows=1, cols=2)
        container.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(container)
        left_cell = container.cell(0, 0)
        right_cell = container.cell(0, 1)

        tou_table = left_cell.add_table(rows=5, cols=2)
        tou_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        h1 = tou_table.cell(0, 0)
        h1.text = ''
        _style_table_cell(h1, is_header=True, bold=True)
        h2 = tou_table.cell(0, 1)
        h2.text = _('Electricity Carbon Dioxide Emissions by Time-Of-Use')
        _style_table_cell(h2, is_header=True, bold=True)
        for i in range(4):
            c1 = tou_table.cell(i + 1, 0)
            c1.text = tou_categories[i]
            _style_table_cell(c1, bold=True)
            c2 = tou_table.cell(i + 1, 1)
            c2.text = str(tou_values[i]) if tou_exists else ''
            _style_table_cell(c2)

        tou_colors = ['#FF1744', '#FF6F00', '#FDD835', '#00BCD4']
        chart_file = None
        if tou_exists and sum(tou_values) > 0:
            chart_file = self._make_pie_chart(tou_values, tou_categories,
                                              _('Electricity Carbon Dioxide Emissions by Time-Of-Use'),
                                              colors=tou_colors)
        if chart_file:
            p = left_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_file, width=Inches(3.2))

        carbon_rows = ca_len + 1
        carbon_table = right_cell.add_table(rows=carbon_rows, cols=2)
        carbon_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hc1 = carbon_table.cell(0, 0)
        hc1.text = ''
        _style_table_cell(hc1, is_header=True, bold=True)
        hc2 = carbon_table.cell(0, 1)
        hc2.text = _('Carbon Dioxide Emissions Proportion')
        _style_table_cell(hc2, is_header=True, bold=True)
        for i in range(ca_len):
            c1 = carbon_table.cell(i + 1, 0)
            c1.text = names[i]
            _style_table_cell(c1, bold=True)
            c2 = carbon_table.cell(i + 1, 1)
            c2.text = str(carbon_values[i]) if carbon_exists else ''
            _style_table_cell(c2)

        carbon_chart_file = None
        if carbon_exists:
            carbon_chart_file = self._make_pie_chart(carbon_values, names,
                                                     _('Carbon Dioxide Emissions Proportion'))
        if carbon_chart_file:
            right_cell.add_paragraph('')
            p = right_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(carbon_chart_file, width=Inches(3.2))

        if self.is_base_period_exists:
            base_period_data = self.report.get('base_period', {})
            bp_names = base_period_data.get('names', [])
            bp_units = base_period_data.get('units', [])
            bp_subtotals = base_period_data.get('subtotals', [])
            bp_subtotals_per_area = base_period_data.get('subtotals_per_unit_area', [])
            bp_increment_rates = base_period_data.get('increment_rates', [])
            bp_total_unit = base_period_data.get('total_unit', 'KGCO2E')
            bp_ca_len = len(bp_names)

            if bp_ca_len > 0:
                doc.add_page_break()

                self._add_heading_styled(doc, self.name + ' - ' +
                                         _('Base Period Carbon Dioxide Emissions'), level=1)

                bp_num_cols = bp_ca_len + 2
                bp_table = doc.add_table(rows=4, cols=bp_num_cols)
                bp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                bp_headers = ['']
                for i in range(bp_ca_len):
                    bp_ui = bp_units[i] if (bp_units and i < len(bp_units)) else ''
                    bp_headers.append(bp_names[i] + ((' (' + bp_ui + ')') if bp_ui else ''))
                bp_headers.append(_('Total') + '(' + bp_total_unit + ')')

                for j, h in enumerate(bp_headers):
                    c = bp_table.cell(0, j)
                    c.text = h
                    _style_table_cell(c, is_header=True, bold=True)

                bp_row_labels = [_('Carbon Dioxide Emissions'), _('Per Unit Area'), _('Increment Rate')]
                for r_idx, lbl in enumerate(bp_row_labels, start=1):
                    c = bp_table.cell(r_idx, 0)
                    c.text = lbl
                    _style_table_cell(c, is_green=True, bold=True)

                bp_total_val = base_period_data.get('total', 0)
                bp_total_per_area = base_period_data.get('total_per_unit_area', None)
                bp_total_inc = base_period_data.get('total_increment_rate', None)

                for i in range(bp_ca_len):
                    col = i + 1
                    c_sub = bp_table.cell(1, col)
                    v = bp_subtotals[i] if (bp_subtotals and i < len(bp_subtotals)) else None
                    c_sub.text = str(round2(v, 2)) if v is not None else ''
                    _style_table_cell(c_sub)

                    c_area = bp_table.cell(2, col)
                    va = bp_subtotals_per_area[i] if (bp_subtotals_per_area and i < len(bp_subtotals_per_area)) else None
                    c_area.text = str(round2(va, 2)) if va is not None else ''
                    _style_table_cell(c_area)

                    c_inc = bp_table.cell(3, col)
                    vi = bp_increment_rates[i] if (bp_increment_rates and i < len(bp_increment_rates)) else None
                    c_inc.text = (str(round2(vi * 100, 2)) + '%') if vi is not None else ''
                    _style_table_cell(c_inc)

                bp_tc = bp_ca_len + 1
                bp_table.cell(1, bp_tc).text = str(round2(bp_total_val, 2))
                _style_table_cell(bp_table.cell(1, bp_tc))
                bp_area_text = str(round2(bp_total_per_area, 2)) if bp_total_per_area is not None else ''
                bp_table.cell(2, bp_tc).text = bp_area_text
                _style_table_cell(bp_table.cell(2, bp_tc))
                bp_inc_text = (str(round2(bp_total_inc * 100, 2)) + '%') if bp_total_inc is not None else ''
                bp_table.cell(3, bp_tc).text = bp_inc_text
                _style_table_cell(bp_table.cell(3, bp_tc))

        doc.add_page_break()

    # ---------- Detailed data ----------
    def _add_detailed_data_section(self, doc):
        _ = self._

        reporting_data = self.report['reporting_period']
        timestamps = reporting_data.get('timestamps', [])

        if not timestamps or len(timestamps[0]) == 0:
            return

        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        values = reporting_data.get('values', [])
        subtotals = reporting_data.get('subtotals', [])
        total_unit = reporting_data.get('total_unit', 'KGCO2E')
        ca_len = len(names)

        rows_per_page = 50

        self._add_heading_styled(doc, self.name + ' ' + _('Detailed Data'), level=1)

        if not self.is_base_period_exists:
            times = timestamps[0]
            if len(times) == 0:
                return

            num_pages = (len(times) + rows_per_page - 1) // rows_per_page

            for page in range(num_pages):
                start_row = page * rows_per_page
                end_row = min(start_row + rows_per_page, len(times))
                page_rows = end_row - start_row

                col_headers = [_('Datetime')]
                for i in range(ca_len):
                    unit_i = units[i] if (units and i < len(units)) else ''
                    col_headers.append(names[i] + ((' (' + unit_i + ')') if unit_i else ''))
                col_headers.append(_('Total') + '(' + total_unit + ')')

                num_cols = len(col_headers)
                table = doc.add_table(rows=page_rows + 2, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for j, h in enumerate(col_headers):
                    c = table.cell(0, j)
                    c.text = h
                    _style_table_cell(c, is_header=True, bold=True, font_size=8)

                for t_idx in range(page_rows):
                    global_idx = start_row + t_idx
                    r_idx = t_idx + 1
                    c0 = table.cell(r_idx, 0)
                    c0.text = str(times[global_idx])
                    _style_table_cell(c0, font_size=8)
                    row_total = 0.0
                    for j in range(ca_len):
                        col = j + 1
                        val = round2(values[j][global_idx], 2) \
                            if j < len(values) and global_idx < len(values[j]) else ''
                        c = table.cell(r_idx, col)
                        c.text = str(val) if val != '' else ''
                        _style_table_cell(c, font_size=8)
                        if j < len(values) and global_idx < len(values[j]):
                            row_total += values[j][global_idx]
                    c_total = table.cell(r_idx, ca_len + 1)
                    c_total.text = str(round2(row_total, 2))
                    _style_table_cell(c_total, font_size=8)

                subtotal_row_idx = page_rows + 1
                c_sub_lbl = table.cell(subtotal_row_idx, 0)
                c_sub_lbl.text = _('Subtotal')
                _style_table_cell(c_sub_lbl, bold=True, font_size=8)
                total_of_subtotals = 0.0
                for i in range(ca_len):
                    col = i + 1
                    c = table.cell(subtotal_row_idx, col)
                    val = subtotals[i] if (subtotals and i < len(subtotals)) else None
                    c.text = str(round2(val, 2)) if val is not None else ''
                    _style_table_cell(c, bold=True, font_size=8)
                    total_of_subtotals += val if val is not None else 0
                c_total_sub = table.cell(subtotal_row_idx, ca_len + 1)
                c_total_sub.text = str(round2(total_of_subtotals, 2))
                _style_table_cell(c_total_sub, bold=True, font_size=8)

                if page < num_pages - 1:
                    doc.add_page_break()
        else:
            base_period_data = self.report['base_period']
            base_timestamps = base_period_data.get('timestamps', [])
            base_values = base_period_data.get('values', [])
            base_subtotals = base_period_data.get('subtotals', [])
            base_names = base_period_data.get('names', [])
            base_units = base_period_data.get('units', [])
            base_ca_len = len(base_names)
            reporting_ca_len = ca_len

            base_times = base_timestamps[0] if base_timestamps else []
            reporting_times = timestamps[0]

            max_len = max(len(base_times), len(reporting_times))
            num_pages = (max_len + rows_per_page - 1) // rows_per_page

            for page in range(num_pages):
                start_row = page * rows_per_page
                end_row = min(start_row + rows_per_page, max_len)
                page_rows = end_row - start_row

                col_headers = [_('Base Period') + ' - ' + _('Datetime')]
                for i in range(base_ca_len):
                    bui = base_units[i] if (base_units and i < len(base_units)) else ''
                    col_headers.append(_('Base Period') + ' - ' + base_names[i] +
                                       ((' (' + bui + ')') if bui else ''))
                col_headers.append(_('Base Period') + ' - ' + _('Total') + '(' + total_unit + ')')
                col_headers.append(_('Reporting Period') + ' - ' + _('Datetime'))
                for i in range(reporting_ca_len):
                    rui = units[i] if (units and i < len(units)) else ''
                    col_headers.append(_('Reporting Period') + ' - ' + names[i] +
                                       ((' (' + rui + ')') if rui else ''))
                col_headers.append(_('Reporting Period') + ' - ' + _('Total') + '(' + total_unit + ')')

                num_cols = len(col_headers)
                table = doc.add_table(rows=page_rows + 2, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for j, h in enumerate(col_headers):
                    c = table.cell(0, j)
                    c.text = h
                    _style_table_cell(c, is_header=True, bold=True, font_size=7)

                for t_idx in range(page_rows):
                    global_idx = start_row + t_idx
                    r_idx = t_idx + 1
                    col = 0
                    c_bt = table.cell(r_idx, col)
                    c_bt.text = base_times[global_idx] if global_idx < len(base_times) else ''
                    _style_table_cell(c_bt, font_size=7)
                    col += 1
                    base_total = 0.0
                    for j in range(base_ca_len):
                        c = table.cell(r_idx, col)
                        if j < len(base_values) and global_idx < len(base_values[j]):
                            c.text = str(round2(base_values[j][global_idx], 2))
                            base_total += base_values[j][global_idx]
                        else:
                            c.text = ''
                        _style_table_cell(c, font_size=7)
                        col += 1
                    c_bt_total = table.cell(r_idx, col)
                    c_bt_total.text = str(round2(base_total, 2)) if global_idx < len(base_times) else ''
                    _style_table_cell(c_bt_total, font_size=7)
                    col += 1
                    c_rt = table.cell(r_idx, col)
                    c_rt.text = reporting_times[global_idx] if global_idx < len(reporting_times) else ''
                    _style_table_cell(c_rt, font_size=7)
                    col += 1
                    rp_total = 0.0
                    for j in range(reporting_ca_len):
                        c = table.cell(r_idx, col)
                        if j < len(values) and global_idx < len(values[j]):
                            c.text = str(round2(values[j][global_idx], 2))
                            rp_total += values[j][global_idx]
                        else:
                            c.text = ''
                        _style_table_cell(c, font_size=7)
                        col += 1
                    c_rt_total = table.cell(r_idx, col)
                    c_rt_total.text = str(round2(rp_total, 2)) if global_idx < len(reporting_times) else ''
                    _style_table_cell(c_rt_total, font_size=7)

                sub_idx = page_rows + 1
                col = 0
                c_s0 = table.cell(sub_idx, col)
                c_s0.text = _('Subtotal')
                _style_table_cell(c_s0, bold=True, font_size=7)
                col += 1
                base_total_all = 0.0
                for i in range(base_ca_len):
                    c = table.cell(sub_idx, col)
                    val = base_subtotals[i] if (base_subtotals and i < len(base_subtotals)) else None
                    c.text = str(round2(val, 2)) if val is not None else ''
                    _style_table_cell(c, bold=True, font_size=7)
                    base_total_all += val if val is not None else 0
                    col += 1
                c_bsub = table.cell(sub_idx, col)
                c_bsub.text = str(round2(base_total_all, 2))
                _style_table_cell(c_bsub, bold=True, font_size=7)
                col += 1
                c_s1 = table.cell(sub_idx, col)
                c_s1.text = _('Subtotal')
                _style_table_cell(c_s1, bold=True, font_size=7)
                col += 1
                rp_total_all = 0.0
                for i in range(reporting_ca_len):
                    c = table.cell(sub_idx, col)
                    val = subtotals[i] if (subtotals and i < len(subtotals)) else None
                    c.text = str(round2(val, 2)) if val is not None else ''
                    _style_table_cell(c, bold=True, font_size=7)
                    rp_total_all += val if val is not None else 0
                    col += 1
                c_rsub = table.cell(sub_idx, col)
                c_rsub.text = str(round2(rp_total_all, 2))
                _style_table_cell(c_rsub, bold=True, font_size=7)

                if page < num_pages - 1:
                    doc.add_page_break()

        doc.add_page_break()

    # ---------- Detailed data charts ----------
    def _add_detailed_data_charts_section(self, doc):
        _ = self._

        reporting_data = self.report['reporting_period']
        timestamps = reporting_data.get('timestamps', [])
        names = reporting_data.get('names', [])
        units = reporting_data.get('units', [])
        values = reporting_data.get('values', [])

        if not timestamps or len(timestamps[0]) == 0 or not names:
            return

        reporting_times = timestamps[0]
        num_categories = len(names)
        charts_per_page = 4

        self._add_heading_styled(doc, self.name + ' ' + _('Detailed Data'), level=1)

        def _safe_list(raw):
            return [(v if isinstance(v, (int, float)) else 0) if v is not None else 0
                    for v in (raw if raw else [])]

        def _set_ticks(ax, raw_len):
            step = max(1, raw_len // 10)
            ax.set_xticks(range(0, raw_len, step))
            ax.set_xticklabels(
                [reporting_times[t][:10] if t < len(reporting_times) else ''
                 for t in range(0, raw_len, step)],
                rotation=45, ha='right', fontsize=7)

        if not self.is_base_period_exists:
            for page_start in range(0, num_categories, charts_per_page):
                page_end = min(page_start + charts_per_page, num_categories)
                page_indices = list(range(page_start, page_end))
                num_on_page = len(page_indices)

                rows = (num_on_page + 1) // 2

                for row_idx in range(rows):
                    slot0 = row_idx * 2
                    slot1 = slot0 + 1
                    has_left = slot0 < num_on_page
                    has_right = slot1 < num_on_page

                    if has_left and not has_right:
                        container = doc.add_table(rows=1, cols=2)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)
                        container.cell(0, 0).merge(container.cell(0, 1))
                        cell = container.cell(0, 0)

                        i = page_indices[slot0]
                        raw_data = values[i] if i < len(values) else []
                        safe_data = _safe_list(raw_data)
                        color = self.chart_colors[i % len(self.chart_colors)]

                        fig, ax = plt.subplots(figsize=(4.8, 3.0))
                        ax.plot(range(len(safe_data)), safe_data, linewidth=1.2,
                                color=color, marker='o', markersize=3,
                                markevery=max(1, len(safe_data) // 30))
                        _set_ticks(ax, len(raw_data))
                        unit_i = units[i] if (units and i < len(units)) else ''
                        ax.set_title(_('Reporting Period Carbon Dioxide Emissions') + ' - ' +
                                     names[i] + ((' (' + unit_i + ')') if unit_i else ''),
                                     fontsize=9, fontweight='bold')
                        ax.grid(True, alpha=0.3)
                        chart_buf = self._fig_to_bytesio(fig, self.dpi)

                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(chart_buf, width=Inches(4.8))
                    else:
                        container_cols = min(2, num_on_page - row_idx * 2)
                        container = doc.add_table(rows=1, cols=container_cols)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)

                        for col_idx in range(container_cols):
                            slot = row_idx * 2 + col_idx
                            if slot >= num_on_page:
                                continue
                            i = page_indices[slot]
                            raw_data = values[i] if i < len(values) else []
                            safe_data = _safe_list(raw_data)
                            color = self.chart_colors[i % len(self.chart_colors)]

                            fig, ax = plt.subplots(figsize=(4.8, 3.0))
                            ax.plot(range(len(safe_data)), safe_data, linewidth=1.2,
                                    color=color, marker='o', markersize=3,
                                    markevery=max(1, len(safe_data) // 30))
                            _set_ticks(ax, len(raw_data))
                            unit_i = units[i] if (units and i < len(units)) else ''
                            ax.set_title(_('Reporting Period Carbon Dioxide Emissions') + ' - ' +
                                         names[i] + ((' (' + unit_i + ')') if unit_i else ''),
                                         fontsize=9, fontweight='bold')
                            ax.grid(True, alpha=0.3)
                            chart_buf = self._fig_to_bytesio(fig, self.dpi)

                            cell = container.cell(0, col_idx)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(chart_buf, width=Inches(4.8))

                if page_end < num_categories:
                    doc.add_page_break()
        else:
            base_period_data = self.report['base_period']
            base_values = base_period_data.get('values', [])
            base_names = base_period_data.get('names', [])

            for page_start in range(0, num_categories, charts_per_page):
                page_end = min(page_start + charts_per_page, num_categories)
                page_indices = list(range(page_start, page_end))
                num_on_page = len(page_indices)

                rows = (num_on_page + 1) // 2

                for row_idx in range(rows):
                    slot0 = row_idx * 2
                    slot1 = slot0 + 1
                    has_left = slot0 < num_on_page
                    has_right = slot1 < num_on_page

                    if has_left and not has_right:
                        container = doc.add_table(rows=1, cols=2)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)
                        container.cell(0, 0).merge(container.cell(0, 1))
                        cell = container.cell(0, 0)

                        i = page_indices[slot0]
                        r_data = values[i] if i < len(values) else []
                        safe_r = _safe_list(r_data)
                        color = self.chart_colors[i % len(self.chart_colors)]

                        fig, ax = plt.subplots(figsize=(4.8, 3.0))
                        ax.plot(range(len(safe_r)), safe_r, linewidth=1.2,
                                color=color, marker='o', markersize=3,
                                markevery=max(1, len(safe_r) // 30),
                                label=_('Reporting Period') + ' - ' + names[i])

                        has_base_line = False
                        if i < len(base_values):
                            b_data = base_values[i]
                            safe_b = _safe_list(b_data)
                            if len(safe_b) > len(safe_r):
                                safe_b = safe_b[:len(safe_r)]
                            x_b = list(range(len(safe_b)))
                            ax.plot(x_b, safe_b, linewidth=1.2,
                                    color=color, linestyle='--', marker='s', markersize=3,
                                    markevery=max(1, len(safe_r) // 30),
                                    label=_('Base Period') + ' - ' +
                                          (base_names[i] if i < len(base_names) else ''))
                            has_base_line = True

                        _set_ticks(ax, len(r_data))
                        unit_i = units[i] if (units and i < len(units)) else ''
                        ax.set_title(
                            _('Base Period Carbon Dioxide Emissions') + ' / ' +
                            _('Reporting Period Carbon Dioxide Emissions') + ' - ' +
                            names[i] + ((' (' + unit_i + ')') if unit_i else ''),
                            fontsize=8, fontweight='bold')
                        if len(safe_r) > 0 or has_base_line:
                            ax.legend(fontsize=7)
                        ax.grid(True, alpha=0.3)
                        chart_buf = self._fig_to_bytesio(fig, self.dpi)

                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(chart_buf, width=Inches(4.8))
                    else:
                        container_cols = min(2, num_on_page - row_idx * 2)
                        container = doc.add_table(rows=1, cols=container_cols)
                        container.alignment = WD_TABLE_ALIGNMENT.CENTER
                        _remove_table_borders(container)

                        for col_idx in range(container_cols):
                            slot = row_idx * 2 + col_idx
                            if slot >= num_on_page:
                                continue
                            i = page_indices[slot]
                            r_data = values[i] if i < len(values) else []
                            safe_r = _safe_list(r_data)
                            color = self.chart_colors[i % len(self.chart_colors)]

                            fig, ax = plt.subplots(figsize=(4.8, 3.0))
                            ax.plot(range(len(safe_r)), safe_r, linewidth=1.2,
                                    color=color, marker='o', markersize=3,
                                    markevery=max(1, len(safe_r) // 30),
                                    label=_('Reporting Period') + ' - ' + names[i])

                            has_base_line = False
                            if i < len(base_values):
                                b_data = base_values[i]
                                safe_b = _safe_list(b_data)
                                if len(safe_b) > len(safe_r):
                                    safe_b = safe_b[:len(safe_r)]
                                x_b = list(range(len(safe_b)))
                                ax.plot(x_b, safe_b, linewidth=1.2,
                                        color=color, linestyle='--', marker='s', markersize=3,
                                        markevery=max(1, len(safe_r) // 30),
                                        label=_('Base Period') + ' - ' +
                                              (base_names[i] if i < len(base_names) else ''))
                                has_base_line = True

                            _set_ticks(ax, len(r_data))
                            unit_i = units[i] if (units and i < len(units)) else ''
                            ax.set_title(
                                _('Base Period Carbon Dioxide Emissions') + ' / ' +
                                _('Reporting Period Carbon Dioxide Emissions') + ' - ' +
                                names[i] + ((' (' + unit_i + ')') if unit_i else ''),
                                fontsize=8, fontweight='bold')
                            if len(safe_r) > 0 or has_base_line:
                                ax.legend(fontsize=6)
                            ax.grid(True, alpha=0.3)
                            chart_buf = self._fig_to_bytesio(fig, self.dpi)

                            cell = container.cell(0, col_idx)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(chart_buf, width=Inches(4.8))

                if page_end < num_categories:
                    doc.add_page_break()

        doc.add_page_break()

    # ---------- Parameters ----------
    def _add_parameters_section(self, doc):
        _ = self._

        params = self.report.get('parameters', {})
        if not params or not params.get('names') or not params.get('timestamps'):
            return

        param_names = params.get('names', [])
        timestamps = params.get('timestamps', [])
        values = params.get('values', [])
        units = params.get('units', [])

        all_zero = True
        for ts_list in timestamps:
            if ts_list and len(ts_list) > 0:
                all_zero = False
                break
        if all_zero:
            return

        all_params = list(range(len(param_names)))
        valid_params = []
        for i in all_params:
            if i < len(timestamps) and len(timestamps[i]) > 0:
                if i < len(values) and len(values[i]) > 0:
                    valid_params.append(i)
        if not valid_params:
            return

        self._add_heading_styled(doc, self.name + ' ' + _('Parameters'), level=1)

        rows_per_param = 10

        for pi in valid_params:
            name = param_names[pi]
            unit_i = units[pi] if (units and pi < len(units)) else ''
            times = timestamps[pi]
            data = values[pi]
            data_len = len(times)

            display_name = name + ((' (' + unit_i + ')') if unit_i else '')

            tbl_rows = min(rows_per_param, data_len)

            fig, ax = plt.subplots(figsize=(5.0, 2.4))
            marker_step_p = max(1, data_len // 20)
            color = '#5B9BD5'
            ax.plot(range(data_len), data, linewidth=1.2,
                    color=color, marker='o', markersize=3,
                    markevery=marker_step_p, label=name)
            ax.fill_between(range(data_len), data, alpha=0.15, color=color)
            step = max(1, data_len // 8)
            ax.set_xticks(range(0, data_len, step))
            ax.set_xticklabels([times[t][:10] for t in range(0, data_len, step)],
                               rotation=45, ha='right', fontsize=6)
            ax.set_ylabel(name, fontsize=8)
            ax.set_title(display_name, fontsize=9, fontweight='bold')
            ax.grid(True, alpha=0.3)
            chart_buf = self._fig_to_bytesio(fig, self.dpi)

            container = doc.add_table(rows=1, cols=2)
            container.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(container)
            left_cell = container.cell(0, 0)
            right_cell = container.cell(0, 1)

            data_table = left_cell.add_table(rows=tbl_rows + 1, cols=2)
            data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            h0 = data_table.cell(0, 0)
            h0.text = _('Time')
            _style_table_cell(h0, is_header=True, bold=True, font_size=8)
            h1 = data_table.cell(0, 1)
            h1.text = name
            _style_table_cell(h1, is_header=True, bold=True, font_size=8)
            for j in range(tbl_rows):
                c_t = data_table.cell(j + 1, 0)
                c_t.text = str(times[j])
                _style_table_cell(c_t, font_size=7)
                c_v = data_table.cell(j + 1, 1)
                c_v.text = str(round2(data[j], 2))
                _style_table_cell(c_v, font_size=7)

            p = right_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_buf, width=Inches(4.5))

    # ---------- Base period existence ----------
    def _is_base_period_timestamp_exists(self, base_period_data: Dict) -> bool:
        timestamps = base_period_data.get('timestamps', [])

        if not timestamps:
            return False

        for timestamp in timestamps:
            if timestamp and len(timestamp) > 0:
                return True

        return False


def export(report,
           name,
           base_period_start_datetime_local,
           base_period_end_datetime_local,
           reporting_start_datetime_local,
           reporting_end_datetime_local,
           period_type,
           language):
    """
    Export report data to DOCX and return base64 encoded string.
    This function maintains the same interface as the Excel exporter.
    """
    exporter = SpaceCarbonDOCXExporter(language)
    return exporter.export(report, name,
                           base_period_start_datetime_local,
                           base_period_end_datetime_local,
                           reporting_start_datetime_local,
                           reporting_end_datetime_local,
                           period_type,
                           language)
